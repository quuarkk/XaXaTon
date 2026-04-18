#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ПДн-сканер корпоративного файлового хранилища (152-ФЗ).

Архитектура:
    1. Инвентаризация   — быстрый обход директорий, сбор списка файлов
    2. Распределение    — файлы разделяются на «лёгкие» (текст/CSV/PDF/HTML…)
                          и «тяжёлые» (изображения/видео — требуют OCR)
    3. Параллельная обработка:
        • ProcessPoolExecutor — для CPU-bound извлечения текста и regex-поиска
        • Отдельный поток / один процесс для OCR (easyocr грузит 1.5 ГБ моделей,
          нельзя клонировать в каждый процесс → OOM на 16 ГБ RAM)
    4. Checkpoint       — результаты стримятся в results.jsonl построчно
                          (падение / перезапуск не теряет прогресс)
    5. Отчёты           — CSV / JSON / Markdown после завершения

Запуск:
    python pdn_scanner.py /path/to/dataset
    python pdn_scanner.py /path/to/dataset --workers 8 --ocr pytesseract
    python pdn_scanner.py /path/to/dataset --resume   # продолжить с чекпоинта
    python pdn_scanner.py /path/to/dataset --no-ocr   # пропустить изображения

Оптимизировано для Ryzen 5 5500U (6c/12t, 16 ГБ RAM).
"""

from __future__ import annotations

import argparse
import csv
import json
import logging
import os
import re
import gzip
import io
import signal
import struct
import sys
import tempfile
import time
import traceback
import zipfile
from collections import defaultdict
from concurrent.futures import ProcessPoolExecutor, ThreadPoolExecutor, as_completed
from dataclasses import asdict, dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Set, Tuple

# ──────────────────────────────────────────────────────────────────────────────
# OCR-БЭКЕНДЫ (EasyOCR на GPU / Tesseract на CPU)
# ──────────────────────────────────────────────────────────────────────────────
#
# Самая тяжёлая часть нашего пайплайна — OCR. В датасете ~2.2 ГБ сканов PDF
# (dataset3 ЮФУ), ~150 TIF (dataset) и видео с документами (dataset4). На CPU
# Tesseract обрабатывает ~0.3 картинки/сек × несколько проходов препроцессинга —
# это часы работы. EasyOCR на NVIDIA GPU даёт 5-10 изображений/сек за один
# проход — ускорение 20-50×.
#
# Стратегия:
#   1. Если есть CUDA и установлен torch+easyocr — используем EasyOCR как
#      дефолт, модель ('ru','en') грузится ОДИН раз на процесс (1.5 ГБ VRAM).
#   2. Иначе откатываемся на pytesseract как сейчас.
#   3. Выбор можно форсировать флагом --ocr-engine.
#
# Важно: EasyOCR модель нельзя клонировать в несколько процессов на одной GPU —
# каждый процесс снова и снова грузит её в VRAM и они толкаются. Поэтому при
# GPU OCR-воркер = 1 штука, а при CPU — сколько скажут.

# Глобал, инициализируется в _init_worker один раз на процесс.
# Формат: ('easyocr', reader_obj) | ('tesseract', None) | None (до init)
_OCR_ENGINE: Optional[Tuple[str, Any]] = None

# Язык для Tesseract (rus+eng) и EasyOCR (['ru','en']) задаётся через CLI.
_OCR_LANG_TESS: str = "rus+eng"
_OCR_LANG_EASY: Tuple[str, ...] = ("ru", "en")

# Размер внутреннего батча EasyOCR — он распознаёт нарезанные текстовые
# регионы пачками. 8-16 хорошо для 6-8 ГБ VRAM, выше — для >= 12 ГБ.
_GPU_BATCH_SIZE: int = 8


def _detect_gpu() -> bool:
    """True, если есть CUDA-GPU. Проверяет torch и paddle (ленивый импорт)."""
    try:
        import torch  # type: ignore
        if torch.cuda.is_available():
            return True
    except Exception:
        pass
    try:
        import paddle  # type: ignore
        if paddle.is_compiled_with_cuda():
            return True
    except Exception:
        pass
    return False


def _has_easyocr() -> bool:
    try:
        import easyocr  # type: ignore  # noqa: F401
        return True
    except Exception:
        return False


def _has_paddleocr() -> bool:
    """PaddleOCR отключён — детектор не находит текст на сканах датасета
    (CUDNN 9.5 vs 9.9 mismatch на RTX 5060 Ti / Blackwell SM_120).
    """
    return False


def _has_surya() -> bool:
    """Surya отключён — зависает при инициализации на Windows + RTX 5060 Ti."""
    return False

def _init_ocr_engine(engine_choice: str = "auto") -> Tuple[str, Any]:
    """Инициализирует OCR-бэкенд ОДИН раз на процесс.

    engine_choice:
        'auto'      → paddleocr (если есть GPU+paddleocr) → easyocr
                      (если есть GPU+easyocr) → tesseract
        'paddleocr' → форсированно PaddleOCR (упадёт если нет пакета)
        'easyocr'   → форсированно EasyOCR
        'tesseract' → форсированно Tesseract
        'none'      → OCR отключён
    Возвращает (engine_name, state_object). state_object для easyocr — это
    Reader, для paddleocr — PaddleOCR объект, для tesseract — None.
    """
    global _OCR_ENGINE
    if _OCR_ENGINE is not None:
        return _OCR_ENGINE

    if engine_choice == "none":
        _OCR_ENGINE = ("none", None)
        return _OCR_ENGINE

    gpu_present = _detect_gpu()

    # ── Приоритет 1: PaddleOCR ───────────────────────────────────────────
    want_paddle = (
        engine_choice == "paddleocr"
        or (engine_choice == "auto" and gpu_present and _has_paddleocr())
    )
    if want_paddle:
        paddle_ocr = _try_init_paddleocr()
        if paddle_ocr is not None:
            _OCR_ENGINE = ("paddleocr", paddle_ocr)
            return _OCR_ENGINE
        # Если пользователь явно просил paddleocr — не откатываемся на easy.
        # Пусть упадёт на tesseract, это даст явный сигнал что-то не так.
        if engine_choice == "paddleocr":
            # дальше пойдём в tesseract fallback
            pass

    # ── Приоритет 2: Surya (PyTorch GPU, 90+ языков) ─────────────────────
    want_surya = (
        engine_choice == "surya"
        or (engine_choice == "auto" and _has_surya())
    )
    if want_surya:
        surya_state = _try_init_surya()
        if surya_state is not None:
            _OCR_ENGINE = ("surya", surya_state)
            return _OCR_ENGINE

    # ── Приоритет 3: EasyOCR ─────────────────────────────────────────────
    want_easy = (
        engine_choice == "easyocr"
        or (engine_choice == "auto" and gpu_present and _has_easyocr())
    )
    if want_easy:
        try:
            import torch  # type: ignore
            import easyocr  # type: ignore
            try:
                torch.backends.cudnn.benchmark = True
            except Exception:
                pass
            gpu_ok = torch.cuda.is_available()
            reader = easyocr.Reader(
                list(_OCR_LANG_EASY),
                gpu=gpu_ok,
                verbose=False,
            )
            _OCR_ENGINE = ("easyocr", reader)
            return _OCR_ENGINE
        except Exception:
            pass

    # ── Приоритет 3: Tesseract (CPU) ──────────────────────────────────────
    try:
        import pytesseract  # type: ignore
        tcmd = os.environ.get("TESSERACT_CMD")
        if tcmd:
            pytesseract.pytesseract.tesseract_cmd = tcmd
        elif os.name == "nt":
            default_win = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
            if Path(default_win).exists():
                pytesseract.pytesseract.tesseract_cmd = default_win
        pytesseract.get_tesseract_version()
        _OCR_ENGINE = ("tesseract", None)
        return _OCR_ENGINE
    except Exception:
        _OCR_ENGINE = ("none", None)
        return _OCR_ENGINE


def _try_init_paddleocr() -> Any:
    """PaddleOCR отключён — несовместимость CUDNN на RTX 5060 Ti (SM_120)."""
    return None


def _try_init_surya():
    """Инициализирует Surya OCR (PyTorch, GPU/CPU, 90+ языков).

    Возвращает dict с {det_predictor, rec_predictor, det_processor, rec_processor}
    или None при любой ошибке.
    """
    try:
        import torch  # type: ignore
        from surya.detection import DetectionPredictor  # type: ignore
        from surya.recognition import RecognitionPredictor  # type: ignore
        device = "cuda" if torch.cuda.is_available() else "cpu"
        det_predictor = DetectionPredictor(device=device)
        rec_predictor = RecognitionPredictor(device=device)
        return {
            "det": det_predictor,
            "rec": rec_predictor,
            "device": device,
        }
    except Exception as e:
        log.debug(f"surya init failed: {e}")
        return None


def _ocr_engine_name() -> str:
    if _OCR_ENGINE is None:
        return "uninit"
    return _OCR_ENGINE[0]


# ──────────────────────────────────────────────────────────────────────────────
# ЛОГИРОВАНИЕ
# ──────────────────────────────────────────────────────────────────────────────

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%H:%M:%S",
    stream=sys.stdout,
    force=True,
)
log = logging.getLogger("pdn_scanner")


# ──────────────────────────────────────────────────────────────────────────────
# РЕГУЛЯРНЫЕ ВЫРАЖЕНИЯ (пересмотренные — меньше ложных срабатываний)
# ──────────────────────────────────────────────────────────────────────────────
#
# Главный принцип: если паттерн опирается только на «длину цифрового поля»
# (СНИЛС = 11 цифр, счёт = 20 цифр, MRZ = 30+ символов ASCII), то без
# **контекстного слова** или **валидации контрольной суммы** он будет
# бесполезен на PDF университета (там тысячи номеров приказов, ИСБН, дат).
# Поэтому каждый такой паттерн требует либо ключевого слова рядом, либо
# валидатора (см. `validate_match`).

# Флаги: IGNORECASE + UNICODE работают по умолчанию у re.compile ниже
_F = re.IGNORECASE | re.UNICODE

PATTERNS: Dict[str, List[re.Pattern]] = {
    # ── ОБЫЧНЫЕ ПДн ──────────────────────────────────────────────────────────
    "Email": [
        re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,24}", _F),
    ],
    "Телефон": [
        # Российские номера: +7/8
        re.compile(
            r"(?<!\d)(?:\+7|8)\s*\(?\d{3}\)?[\s\-]?\d{3}[\s\-]?\d{2}[\s\-]?\d{2}(?!\d)",
            _F,
        ),
        # Международные номера: +[1-3 цифры кода страны] + 7-12 цифр
        re.compile(
            r"(?<!\d)\+(?!7(?!\d))(?:[2-9]\d{0,2})[\s\-]?\(?\d{1,4}\)?[\s\-]?\d{2,5}[\s\-]?\d{2,5}(?:[\s\-]?\d{1,4})?(?!\d)",
            _F,
        ),
    ],
    "Дата_рождения": [
        # Дата рождения с контекстом — иначе слишком шумно
        re.compile(
            r"(?:дата\s+рожд\w*|рождени[ея]|д\.р\.|date\s+of\s+birth|datum\s+narozeni|fecha\s+de\s+nacimiento|geburts\w*|né\s*le|data\s+di\s+nascita|dob)[\s:]*"
            r"((?:0[1-9]|[12]\d|3[01])[.\/\-](?:0[1-9]|1[0-2])[.\/\-](?:19|20)\d{2})",
            _F,
        ),
        # Место рождения — ТЗ требует «Дата и место рождения» как одну категорию.
        # Требуем ключевое слово, чтобы не ловить любые упоминания городов.
        re.compile(
            r"(?:место\s+рожд\w*|м\.?\s*р\.?|place\s+of\s+birth|birthplace)[\s.:]*"
            r"([А-ЯЁA-Za-z][а-яёА-ЯЁA-Za-z\s\-\.]{2,60})",
            _F,
        ),
        # «Родился(-ась) в г./с./п. <название>»
        re.compile(
            r"(?:родил(?:ся|ась|и)\s+в\s+)(?:г\.|город|с\.|село|п\.|посёлок|пос\.)?\s*"
            r"([А-ЯЁA-Za-z][а-яёА-ЯЁA-Za-z\s\-]{2,40})",
            _F,
        ),
    ],
    "ФИО": [
        # Надёжнее всего — имя + отчество с характерными суффиксами
        re.compile(
            r"\b[А-ЯЁ][а-яё]{1,20}\s+"
            r"[А-ЯЁ][а-яё]{1,20}(?:ович|евич|ич|овна|евна|инична|ична)\b",
            _F,
        ),
        # Или «Фамилия И.О.» / «Фамилия И.»
        re.compile(
            r"\b[А-ЯЁ][а-яё]{2,20}\s+[А-ЯЁ]\.\s*[А-ЯЁ]\.", _F,
        ),
        # Или с явной меткой «ФИО: …»
        re.compile(
            r"(?:ФИО|Ф\.И\.О\.)[\s:]+"
            r"([А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+(?:\s+[А-ЯЁ][а-яё]+)?)",
            _F,
        ),
        # Поля имени в международных документах / ID-картах
        re.compile(
            r"(?:surname|last\s+name|family\s+name|given\s+name|first\s+name|"
            r"name|jmeno|prijmeni)[\s:]+"
            r"([A-Z][A-Za-z'\-]{1,30}(?:\s+[A-Z][A-Za-z'\-]{1,30}){0,2})",
            _F,
        ),
    ],
    "Адрес": [
        # Город + улица + дом — три якорных элемента
        re.compile(
            r"(?:г\.|город)\s*[А-ЯЁ][а-яё\-]{2,25}[,;]?\s*"
            r"(?:ул\.|улица|пр\.|проспект|пер\.|переулок|б-р|бульвар|ш\.|шоссе)\s*"
            r"[А-ЯЁа-яё\d\s\-\.]{2,40}[,;]?\s*"
            r"(?:д\.|дом)\s*\d+[а-яА-Я]?",
            _F,
        ),
        # Индекс + город
        re.compile(r"\b\d{6}[,\s]+(?:г\.|город)\s*[А-ЯЁ][а-яё\-]{2,25}", _F),
    ],

    # ── ГОСУДАРСТВЕННЫЕ ИДЕНТИФИКАТОРЫ ───────────────────────────────────────
    "Паспорт_РФ": [
        # С явной меткой «паспорт / серия»
        re.compile(
            r"(?:паспорт|серия\s*(?:и\s*номер)?|passport)[\s№:]*"
            r"(\d{2}\s?\d{2}\s?№?\s?\d{6})",
            _F,
        ),
        # Классический формат «4503 123456» (две пары цифр + 6 цифр через пробел)
        # — без групп захвата, чтобы m.group(0) возвращал полный номер
        re.compile(r"(?<!\d)\d{4}\s\d{6}(?!\d)"),
        # Международные удостоверения личности: document/id/card number
        re.compile(
            r"(?:document\s*(?:no|number)|id\s*(?:no|number)|card\s*(?:no|number)|"
            r"licen[cs]e\s*(?:no|number)|obcansky\s+prukaz)[\s№:#-]*"
            r"([A-Z0-9]{6,12})",
            _F,
        ),
    ],
    "СНИЛС": [
        # С разделителями — самая надёжная форма
        re.compile(r"\b\d{3}[\s\-]\d{3}[\s\-]\d{3}[\s\-]\d{2}\b"),
        # Без разделителей — только с ключевым словом
        re.compile(r"(?:СНИЛС|страховой\s+номер)[\s№:]*(\d{11})", _F),
    ],
    "ИНН": [
        # С ключевым словом
        re.compile(r"(?:ИНН|INN)[\s№:]*(\d{10}|\d{12})", _F),
    ],
    "Водительское_удостоверение": [
        # С ключевым словом
        re.compile(
            r"(?:водит(?:ельское)?\s+удостов\w*|ВУ|driver\s+licen[cs]e)"
            r"[\s№:]*(\d{2}\s?[А-ЯA-Z]{2}\s?\d{6})",
            _F,
        ),
    ],
    "Идентификационный_документ": [
        # Общий эвристический маркер для зарубежных ID-карт / удостоверений.
        # Сам по себе keyword ещё не означает ПДн, поэтому ниже в CONTEXT_REQUIRED
        # требуем рядом дату или документный номер.
        re.compile(
            r"(?:passport|identification\s+card|identity\s+card|national\s+id|"
            r"id\s+card|residence\s+permit|obcansk\w*\s+pruk\w*)",
            _F,
        ),
    ],
    "MRZ": [
        # Двухстрочная MRZ TD3 (паспорт): 44 + 44 символа.
        # ВАЖНО: в настоящей MRZ обязательны символы-заполнители '<' и буквы
        # (коды стран RUS/USA, фамилия). Без них это просто цифровая строка
        # (номера приказов, табличные данные). Требуем наличие '<' или буквы.
        re.compile(
            r"(?=[A-Z0-9<]*[A-Z<])[A-Z0-9<]{44}\s*\n\s*[A-Z0-9<]{44}"
        ),
        # TD1 (ID): 30+30+30
        re.compile(
            r"(?=[A-Z0-9<]*[A-Z<])[A-Z0-9<]{30}\s*\n\s*[A-Z0-9<]{30}\s*\n\s*[A-Z0-9<]{30}"
        ),
    ],

    # ── ПЛАТЁЖНАЯ ИНФОРМАЦИЯ ─────────────────────────────────────────────────
    "Банковская_карта": [
        # Любая последовательность 13-19 цифр с опц. разделителями;
        # обязательно пройдёт алгоритм Луна в validate_match().
        re.compile(r"(?<!\d)(?:\d{4}[\s\-]?){3}\d{1,4}(?:[\s\-]?\d{3})?(?!\d)"),
    ],
    "CVV": [
        # Только с ключевым словом — иначе бесполезно
        re.compile(r"\bCVV2?\b[\s:№]*(\d{3,4})\b", _F),
        re.compile(r"\bCVC\b[\s:№]*(\d{3,4})\b", _F),
    ],
    "БИК": [
        # БИК всегда с ключевым словом
        re.compile(r"\bБИК\b[\s:№]*(\d{9})\b", _F),
    ],
    "Банковский_счёт": [
        # Только с меткой — 20 цифр в PDF = номер приказа, ISBN, телефон и т.п.
        re.compile(
            r"(?:р\.?\s*с\.?|р/с|расч[её]тн\w*\s+сч[её]т|"
            r"лицев\w*\s+сч[её]т|счет|сч[её]т)[\s№:.]*(\d{20})\b",
            _F,
        ),
    ],

    # ── БИОМЕТРИЯ ────────────────────────────────────────────────────────────
    "Биометрия": [
        re.compile(
            r"(?:отпечат(?:ок|ки)\s+пальц|дактилоскоп\w+|"
            r"радужн\w+\s+оболочк\w+|сетчатк\w+\s+глаз\w*|"
            r"голосов\w+\s+(?:слепок|образ(?:ец|цы))|геометри\w+\s+лиц\w+|"
            r"биометрическ\w+\s+(?:данн|идентиф|образ)|"
            r"fingerprint|iris\s+scan|face\s+recognition|voice\s+biometric)",
            _F,
        ),
    ],

    # ── СПЕЦ. КАТЕГОРИИ ──────────────────────────────────────────────────────
    "Здоровье": [
        # Диагноз с содержанием (не просто одинокое слово)
        re.compile(
            r"(?:диагноз|анамнез)[\s:]+[А-Яа-яA-Za-z0-9\s,.\-]{5,120}", _F,
        ),
        # МЕДИЦИНСКОЕ заключение (исключаем юридическое «заключение договора»).
        # Требуем контекста «медицинское / врачебное / клиническое заключение».
        re.compile(
            r"(?:медицинск\w+|врачебн\w+|клиническ\w+|психиатрическ\w+)"
            r"\s+заключени[ея]",
            _F,
        ),
        # МКБ-10 коды — очень специфичны
        re.compile(r"\bМКБ[-\s]?10\b[\s:№]*[A-ZА-Я]\d{2}(?:\.\d{1,2})?", _F),
        re.compile(r"\bICD[-\s]?10\b[\s:№]*[A-Z]\d{2}(?:\.\d{1,2})?", _F),
        # Явные упоминания реальных заболеваний / учёта
        re.compile(
            r"(?:история\s+болезни|психиатрическ\w+\s+диспансер|"
            r"наркологическ\w+\s+(?:диспансер|учёт|учет)|"
            r"онкологическ\w+\s+(?:заболевани|диагноз)|"
            r"ВИЧ[-\s]?(?:инфекц|позитивн)|гепатит\s+[BC]|"
            r"инвалидност[ьи]\s+\d+\s+групп|"
            r"листок\s+нетрудоспособност)",
            _F,
        ),
    ],
    "Религия": [
        re.compile(
            r"(?:вероисповедани[ея]|религиозн\w+\s+(?:принадлежност|убеждени))"
            r"[\s:]+[А-Яа-я\s]{3,40}",
            _F,
        ),
    ],
    "Политика": [
        re.compile(
            r"(?:политическ\w+\s+убеждени|член\s+партии\s+[«\"А-Яа-я])",
            _F,
        ),
    ],
    "Национальность_раса": [
        re.compile(
            r"(?:национальност[ьи]|этническ\w+\s+происхождени|расов\w+\s+принадлежност)"
            r"[\s:]+[А-Яа-я\s]{3,40}",
            _F,
        ),
    ],
}

# Категории ПДн → уровень защищённости
CATEGORIES_UZ1 = {"Биометрия", "Здоровье", "Религия", "Политика", "Национальность_раса"}
CATEGORIES_UZ2_STRONG = {"Банковская_карта", "CVV", "БИК", "Банковский_счёт"}
CATEGORIES_GOV_ID = {"Паспорт_РФ", "СНИЛС", "ИНН",
                     "Водительское_удостоверение", "Идентификационный_документ", "MRZ"}
CATEGORIES_BASIC = {"ФИО", "Email", "Телефон", "Адрес", "Дата_рождения"}


# ──────────────────────────────────────────────────────────────────────────────
# ВАЛИДАТОРЫ
# ──────────────────────────────────────────────────────────────────────────────

def luhn_check(number: str) -> bool:
    """
    Алгоритм Луна для номеров банковских карт + BIN-фильтр.

    Алгоритм Луна имеет ~10% случайных срабатываний на произвольных цифровых
    строках. В юридических документах с номерами приказов/статей это даёт
    десятки ложных "карт". Усиливаем двумя проверками:

    1. Длина номера — только стандартные: 13 (Visa), 15 (Amex), 16 (большинство),
       18 (редкие), 19 (Maestro, UnionPay). Длины 14, 17 не используются.
    2. BIN-префикс — первая цифра должна быть 2–6:
         2  — Mir, Mastercard (2221-2720), UnionPay
         3  — Amex (34/37), JCB, Diners
         4  — Visa
         5  — Mastercard, Maestro
         6  — Discover, Maestro, UnionPay
       Номера, начинающиеся на 0, 1, 7, 8, 9 — не банковские карты.
    """
    digits = [int(d) for d in re.sub(r"\D", "", number)]
    if len(digits) not in (13, 15, 16, 18, 19):
        return False
    if digits[0] not in (2, 3, 4, 5, 6):
        return False
    checksum = 0
    for i, d in enumerate(reversed(digits)):
        if i % 2 == 1:
            d *= 2
            if d > 9:
                d -= 9
        checksum += d
    return checksum % 10 == 0


def snils_check(snils: str) -> bool:
    digits = re.sub(r"\D", "", snils)
    if len(digits) != 11:
        return False
    n = [int(d) for d in digits]
    ctrl = n[9] * 10 + n[10]
    s = sum((9 - i) * n[i] for i in range(9))
    if s > 101:
        s %= 101
    if s in (100, 101):
        s = 0
    return s == ctrl


def inn_check(inn: str) -> bool:
    digits = re.sub(r"\D", "", inn)
    if len(digits) == 10:
        w = [2, 4, 10, 3, 5, 9, 4, 6, 8]
        c = sum(wi * int(di) for wi, di in zip(w, digits[:9])) % 11 % 10
        return c == int(digits[9])
    if len(digits) == 12:
        w1 = [7, 2, 4, 10, 3, 5, 9, 4, 6, 8]
        w2 = [3, 7, 2, 4, 10, 3, 5, 9, 4, 6, 8]
        c1 = sum(wi * int(di) for wi, di in zip(w1, digits[:10])) % 11 % 10
        c2 = sum(wi * int(di) for wi, di in zip(w2, digits[:11])) % 11 % 10
        return c1 == int(digits[10]) and c2 == int(digits[11])
    return False


def validate_match(category: str, raw: str) -> bool:
    """Дополнительная валидация. True = считаем находку валидной."""
    clean = re.sub(r"\D", "", raw)
    if category == "Банковская_карта":
        return luhn_check(raw)
    if category == "СНИЛС":
        # Если это 11 цифр подряд без разделителей — обязателен контроль
        if len(clean) == 11 and not any(c in raw for c in " -"):
            return snils_check(raw)
        return len(clean) == 11 and snils_check(raw)
    if category == "ИНН":
        return inn_check(raw)
    if category == "Банковский_счёт":
        return len(clean) == 20
    if category == "БИК":
        # БИК РФ: начинается на 04 (Россия)
        return len(clean) == 9 and clean.startswith("04")
    if category == "Паспорт_РФ":
        return len(clean) == 10
    return True


def mask(value: str) -> str:
    """Маскирование значения ПДн для отчёта."""
    s = value.strip()
    if len(s) <= 4:
        return "*" * len(s)
    vis = max(2, len(s) // 5)
    return s[:vis] + "*" * (len(s) - 2 * vis) + s[-vis:]


# ──────────────────────────────────────────────────────────────────────────────
# ПОИСК ПДн В ТЕКСТЕ
# ──────────────────────────────────────────────────────────────────────────────

def detect_pdn(text: str, max_findings_per_cat: int = 200) -> Dict[str, List[str]]:
    """
    Ищет все категории ПДн в тексте.
    Возвращает словарь {категория: [маскированные значения]}.

    max_findings_per_cat: лимит находок в одной категории на файл —
    защита от раздувания отчёта на больших CSV (100к записей).
    """
    if not text or len(text) < 5:
        return {}

    found: Dict[str, List[str]] = defaultdict(list)
    seen: Set[Tuple[str, str]] = set()

    for category, patterns in PATTERNS.items():
        for pattern in patterns:
            try:
                for m in pattern.finditer(text):
                    if len(found[category]) >= max_findings_per_cat:
                        break
                    # Если есть группа захвата — берём именно её (значение без ключевого
                    # слова), иначе всё совпадение. Это позволяет корректно дедуплицировать
                    # «Паспорт 4503 123456» и «4503 123456» как одно ПДн.
                    raw_full = m.group(0).strip()
                    if m.lastindex:
                        raw = m.group(m.lastindex).strip()
                    else:
                        raw = raw_full
                    key = (category, re.sub(r"\s+", "", raw).lower())
                    if key in seen:
                        continue
                    seen.add(key)
                    # Валидация проводится по чистому значению
                    if validate_match(category, raw):
                        found[category].append(mask(raw))
            except re.error:
                continue

    # ВАЖНО: удаляем пустые категории. Если regex сматчил кандидатов, но все
    # отфильтровала validate_match — категория остаётся в defaultdict с пустым
    # списком. Тогда classify_uz(set(found.keys())) ошибочно примет её как
    # найденную и завысит УЗ. Чистим.
    return {k: v for k, v in found.items() if v}


def classify_uz(categories: Set[str], total_findings: int) -> str:
    """Классификация уровня защищённости по 152-ФЗ."""
    if not categories:
        return "—"
    if categories & CATEGORIES_UZ1:
        return "УЗ-1"
    if categories & CATEGORIES_UZ2_STRONG:
        return "УЗ-2"
    gov = categories & CATEGORIES_GOV_ID
    if gov:
        # Много гос. ID → УЗ-2, иначе УЗ-3
        return "УЗ-2" if total_findings > 100 else "УЗ-3"
    if categories & CATEGORIES_BASIC:
        return "УЗ-3" if total_findings > 50 else "УЗ-4"
    return "—"


# ──────────────────────────────────────────────────────────────────────────────
# КОНТЕКСТНЫЕ ФИЛЬТРЫ
# ──────────────────────────────────────────────────────────────────────────────
#
# Regex + алгоритм Луна вдвоём дают ~1% ложных срабатываний на длинных
# цифровых потоках (номера приказов, ISBN, табличные данные). В документах
# без ПЛАТЁЖНОЙ тематики этот 1% превращается в десятки «карт». Фильтр:
# если в тексте файла нет ни одного слова, указывающего на платёжный контекст —
# отбрасываем найденные «карты». Аналогично для некоторых других категорий.

CONTEXT_REQUIRED: Dict[str, re.Pattern] = {
    "Банковская_карта": re.compile(
        r"(?:карт[аыуеой]|card|visa|master|master\s*card|мир|maestro|"
        r"платеж|платёж|оплат|cvc|cvv|номер\s+счёт|номер\s+счет|"
        r"банк\w*\s+счёт|банк\w*\s+счет|реквизит|transaction|эквайринг)",
        re.IGNORECASE,
    ),
    "Идентификационный_документ": re.compile(
        r"(?:(?:0?[1-9]|[12]\d|3[01])[./-](?:0?[1-9]|1[0-2])[./-](?:\d{2}|\d{4})|"
        r"[A-Z0-9]{6,12})",
        re.IGNORECASE,
    ),
}


def apply_context_filter(text: str, found: Dict[str, List[str]]) -> Dict[str, List[str]]:
    """Убирает категории, для которых в тексте нет обязательного контекста."""
    out = {}
    for cat, vals in found.items():
        req = CONTEXT_REQUIRED.get(cat)
        if req and not req.search(text):
            continue  # контекста нет → отбрасываем целиком
        out[cat] = vals
    return out


# ──────────────────────────────────────────────────────────────────────────────
# ДЕТЕКЦИЯ TESSERACT (OCR-движок)
# ──────────────────────────────────────────────────────────────────────────────

def tesseract_available() -> bool:
    """
    Проверяет, доступен ли Tesseract. Вызывается из main() до запуска пулов.
    """
    try:
        import pytesseract
        tcmd = os.environ.get("TESSERACT_CMD")
        if tcmd:
            pytesseract.pytesseract.tesseract_cmd = tcmd
        elif os.name == "nt":
            default_win = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
            if Path(default_win).exists():
                pytesseract.pytesseract.tesseract_cmd = default_win
        # Это бросит TesseractNotFoundError, если бинарь не найден
        pytesseract.get_tesseract_version()
        return True
    except Exception:
        return False


# ──────────────────────────────────────────────────────────────────────────────
# ЭКСТРАКТОРЫ ТЕКСТА
# ──────────────────────────────────────────────────────────────────────────────

# На один файл принимаем максимум 8 МБ текста для regex.
# Это ~50 тыс. строк CSV — достаточно, чтобы найти ПДн, и не ложит память.
MAX_TEXT_BYTES = 8 * 1024 * 1024
# Для структурированных данных: анализируем чанками по N строк.
CHUNK_ROWS = 20_000
# Пропускаем изображения > 25 МБ (OCR будет работать минутами)
MAX_IMAGE_BYTES = 25 * 1024 * 1024
# Hard-таймаут на файл (сек)
FILE_TIMEOUT_SEC = 60   # таймаут на файл (сек); OCR-файлы не должны висеть дольше


def _detect_encoding(raw: bytes) -> str:
    """Пытаемся определить кодировку. Fallback — utf-8 с заменой ошибок."""
    try:
        import chardet
        enc = chardet.detect(raw[:64_000]).get("encoding")
        if enc:
            return enc
    except Exception:
        pass
    # Типичные русские кодировки
    for enc in ("utf-8", "cp1251", "koi8-r", "utf-16"):
        try:
            raw[:10_000].decode(enc)
            return enc
        except UnicodeDecodeError:
            continue
    return "utf-8"


def extract_csv(path: Path) -> Dict[str, List[str]]:
    """
    CSV: построчное чтение с чанками. Собираем находки сразу, не накапливая
    гигантский текст в памяти.
    """
    with open(path, "rb") as f:
        raw = f.read(64_000)
    enc = _detect_encoding(raw)

    aggregate: Dict[str, List[str]] = defaultdict(list)
    seen: Set[Tuple[str, str]] = set()

    buf: List[str] = []
    row_count = 0
    try:
        with open(path, "r", encoding=enc, errors="replace", newline="") as f:
            # Пробуем определить диалект
            sample = f.read(4096)
            f.seek(0)
            try:
                dialect = csv.Sniffer().sniff(sample, delimiters=",;|\t")
            except csv.Error:
                dialect = csv.excel
            reader = csv.reader(f, dialect)
            for row in reader:
                buf.append(" | ".join(str(c) for c in row))
                row_count += 1
                if row_count % CHUNK_ROWS == 0:
                    _merge_findings(aggregate, seen, "\n".join(buf))
                    buf.clear()
        if buf:
            _merge_findings(aggregate, seen, "\n".join(buf))
    except Exception as e:
        log.debug(f"CSV read fallback {path.name}: {e}")
        # Fallback — читаем как обычный текст
        try:
            with open(path, "r", encoding=enc, errors="replace") as f:
                _merge_findings(aggregate, seen, f.read(MAX_TEXT_BYTES))
        except Exception:
            pass

    return dict(aggregate)


def extract_json(path: Path) -> Dict[str, List[str]]:
    """JSON: при размере > 10 МБ используем построчный стриминг (ijson)."""
    aggregate: Dict[str, List[str]] = defaultdict(list)
    seen: Set[Tuple[str, str]] = set()

    size = path.stat().st_size
    if size < 10 * 1024 * 1024:
        try:
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                data = json.load(f)
            text = json.dumps(data, ensure_ascii=False)
            _merge_findings(aggregate, seen, text[:MAX_TEXT_BYTES])
            return dict(aggregate)
        except Exception:
            # Может быть JSONL или битый JSON — падём через fallback
            pass

    # Большой / битый JSON → читаем как текст построчно
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            buf: List[str] = []
            total = 0
            for line in f:
                buf.append(line)
                total += len(line)
                if total > 2 * 1024 * 1024:  # чанк 2 МБ
                    _merge_findings(aggregate, seen, "".join(buf))
                    buf.clear()
                    total = 0
            if buf:
                _merge_findings(aggregate, seen, "".join(buf))
    except Exception as e:
        log.debug(f"JSON fallback fail {path.name}: {e}")
    return dict(aggregate)


def extract_parquet(path: Path) -> Dict[str, List[str]]:
    """Parquet: читаем через pyarrow по row groups для экономии памяти."""
    aggregate: Dict[str, List[str]] = defaultdict(list)
    seen: Set[Tuple[str, str]] = set()
    try:
        import pyarrow.parquet as pq
        pf = pq.ParquetFile(str(path))
        for rg_idx in range(pf.num_row_groups):
            table = pf.read_row_group(rg_idx)
            # Конвертим в CSV-строки колонок
            df = table.to_pandas()
            text = df.astype(str).apply(
                lambda r: " | ".join(r.values), axis=1
            ).str.cat(sep="\n")
            _merge_findings(aggregate, seen, text)
    except Exception as e:
        log.debug(f"Parquet read {path.name}: {e}")
        # Fallback через pandas
        try:
            import pandas as pd
            df = pd.read_parquet(path)
            _merge_findings(aggregate, seen, df.to_csv(index=False)[:MAX_TEXT_BYTES])
        except Exception as e2:
            log.debug(f"Parquet pandas fallback fail {path.name}: {e2}")
    return dict(aggregate)


def extract_pdf(path: Path) -> str:
    """
    PDF: сначала быстрый PyMuPDF (fitz), при его отсутствии — pdfplumber, затем
    pdfminer / PyPDF2 как запасные варианты.

    После извлечения основного текста дополнительно сканируем:
      • /EmbeddedFiles — файловые вложения (Word, Excel, изображения)

    ВАЖНО: если fitz успешно ОТКРЫЛ файл (без исключения), но не смог извлечь текст —
    это почти всегда означает, что PDF представляет собой скан без text-слоя.
    Fallback-библиотеки (pdfplumber/pypdf) в этом случае тоже ничего не найдут,
    но потратят секунды и засорят лог warnings'ами. Поэтому не идём в fallback,
    а сразу возвращаем пустоту. Для таких PDF нужен OCR-пайплайн.
    """
    main_text = ""
    fitz_ok = False

    # 1. PyMuPDF (самый быстрый)
    try:
        import fitz  # PyMuPDF
        parts = []
        with fitz.open(str(path)) as doc:
            fitz_ok = True  # файл открылся без исключения
            # Ограничим: первые 100 страниц хватит, чтобы понять, есть ли ПДн
            for i, page in enumerate(doc):
                if i >= 100:
                    break
                t = page.get_text("text")
                if t:
                    parts.append(t)
                if sum(len(p) for p in parts) > MAX_TEXT_BYTES:
                    break
        main_text = "\n".join(parts)
        if not main_text.strip() and fitz_ok:
            # fitz открыл файл, но текст пустой → это скан. Fallback бесполезен.
            pass
    except ImportError:
        pass
    except Exception as e:
        log.debug(f"fitz fail {path.name}: {e}")

    # Fallback-цепочка — только если fitz вообще не смог открыть файл
    if not main_text.strip() and not fitz_ok:
        # 2. pdfplumber
        try:
            import pdfplumber
            parts = []
            with pdfplumber.open(str(path)) as pdf:
                for i, page in enumerate(pdf.pages[:50]):
                    t = page.extract_text()
                    if t:
                        parts.append(t)
            main_text = "\n".join(parts)
        except ImportError:
            pass
        except Exception as e:
            log.debug(f"pdfplumber fail {path.name}: {e}")

    if not main_text.strip() and not fitz_ok:
        # 3. pypdf / PyPDF2
        for modname in ("pypdf", "PyPDF2"):
            try:
                mod = __import__(modname)
                with open(path, "rb") as f:
                    reader = mod.PdfReader(f)
                    parts = []
                    for i, page in enumerate(reader.pages):
                        if i >= 50:
                            break
                        t = page.extract_text() or ""
                        parts.append(t)
                main_text = "\n".join(parts)
                if main_text.strip():
                    break
            except ImportError:
                continue
            except Exception:
                continue

    # 4. Вложения (/EmbeddedFiles) — сканируем всегда, независимо от основного текста
    attachments_text = _extract_pdf_attachments(path)

    combined = "\n".join(filter(None, [main_text, attachments_text]))
    return combined[:MAX_TEXT_BYTES]


def extract_pdf_ocr(path: Path, max_pages: int = 20) -> str:
    """OCR-резерв для сканированных PDF без text-layer.

    Рендерит только страницы, на которых текстовый слой пустой или почти пустой.
    Это закрывает главный пробел исходной версии: PDF-сканы раньше честно
    считались "без ПДн", хотя фактически содержали изображение документа.

    На EasyOCR (GPU) рендер PDF-страницы — главное CPU-узкое место, поэтому
    понижена матрица upscale (2.0x вместо 2.5x): детектору EasyOCR этого
    достаточно, а rendering ускоряется ~30%.
    """
    try:
        import fitz
        from PIL import Image
    except ImportError:
        return ""

    if _OCR_ENGINE is None:
        _init_ocr_engine("auto")
    if _ocr_engine_name() == "none":
        return OCR_UNAVAILABLE_SENTINEL

    # На tesseract детектору нужна чуть большая детализация — поднимаем матрицу.
    zoom = 2.0 if _ocr_engine_name() == "easyocr" else 2.5

    try:
        texts = []
        with fitz.open(str(path)) as doc:
            for i, page in enumerate(doc):
                if i >= max_pages:
                    break
                layer_text = page.get_text("text") or ""
                compact = re.sub(r"\s+", "", layer_text)
                if len(compact) >= 40:
                    continue

                pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
                mode = "RGB" if pix.n < 4 else "RGBA"
                img = Image.frombytes(mode, [pix.width, pix.height], pix.samples)
                if mode == "RGBA":
                    img = img.convert("RGB")

                t = _ocr_pil_image(img)
                if t == OCR_UNAVAILABLE_SENTINEL:
                    return OCR_UNAVAILABLE_SENTINEL
                if t and t.strip():
                    texts.append(t)

        return "\n".join(texts)[:MAX_TEXT_BYTES]
    except Exception as e:
        log.debug(f"pdf ocr {path.name}: {e}")
        return ""


def extract_docx(path: Path) -> str:
    parts = []
    try:
        import docx
        doc = docx.Document(str(path))
        for p in doc.paragraphs:
            if p.text:
                parts.append(p.text)
        # И таблицы
        for tbl in doc.tables:
            for row in tbl.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
    except Exception as e:
        log.debug(f"docx {path.name}: {e}")

    main_text = "\n".join(parts)

    # Дополнительно: embedded OLE-объекты (вставки из других документов)
    embedded_text = _extract_docx_embedded(path)

    return "\n".join(filter(None, [main_text, embedded_text]))[:MAX_TEXT_BYTES]


def extract_doc_rtf(path: Path) -> str:
    """DOC/RTF: docx2txt справляется со многими .doc через встроенный антиворд."""
    try:
        import docx2txt
        return (docx2txt.process(str(path)) or "")[:MAX_TEXT_BYTES]
    except Exception:
        pass
    # Fallback: striprtf для rtf
    if path.suffix.lower() == ".rtf":
        try:
            from striprtf.striprtf import rtf_to_text
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                return rtf_to_text(f.read())[:MAX_TEXT_BYTES]
        except Exception:
            pass
    # Грубый fallback: байты с заменой
    try:
        raw = path.read_bytes()[:MAX_TEXT_BYTES]
        return raw.decode("utf-8", errors="replace")
    except Exception:
        return ""


def extract_xls(path: Path) -> Dict[str, List[str]]:
    """XLS/XLSX: стримим через pandas по листам, чанками."""
    aggregate: Dict[str, List[str]] = defaultdict(list)
    seen: Set[Tuple[str, str]] = set()
    try:
        import pandas as pd
        # openpyxl для xlsx, xlrd для старого xls
        engine = "openpyxl" if path.suffix.lower() == ".xlsx" else None
        sheets = pd.read_excel(path, sheet_name=None, dtype=str, engine=engine)
        for name, df in sheets.items():
            if df is None or df.empty:
                continue
            # Конкатим по строкам и скармливаем чанками
            for start in range(0, len(df), CHUNK_ROWS):
                chunk = df.iloc[start:start + CHUNK_ROWS]
                text = chunk.astype(str).apply(
                    lambda r: " | ".join(r.values), axis=1
                ).str.cat(sep="\n")
                _merge_findings(aggregate, seen, text)
    except Exception as e:
        log.debug(f"xls {path.name}: {e}")
    return dict(aggregate)


def extract_html(path: Path) -> str:
    try:
        import bs4
        raw = path.read_bytes()
        enc = _detect_encoding(raw)
        soup = bs4.BeautifulSoup(
            raw.decode(enc, errors="replace"),
            "html.parser",
        )
        for tag in soup(["script", "style"]):
            tag.decompose()
        return soup.get_text(separator=" ", strip=True)[:MAX_TEXT_BYTES]
    except ImportError:
        # Fallback: грубая очистка
        try:
            raw = path.read_bytes()
            enc = _detect_encoding(raw)
            txt = raw.decode(enc, errors="replace")
            txt = re.sub(r"<script[^>]*>.*?</script>", " ", txt,
                         flags=re.DOTALL | re.IGNORECASE)
            txt = re.sub(r"<style[^>]*>.*?</style>", " ", txt,
                         flags=re.DOTALL | re.IGNORECASE)
            return re.sub(r"<[^>]+>", " ", txt)[:MAX_TEXT_BYTES]
        except Exception:
            return ""
    except Exception as e:
        log.debug(f"html {path.name}: {e}")
        return ""


def extract_txt(path: Path) -> str:
    try:
        raw = path.read_bytes()[:MAX_TEXT_BYTES]
        enc = _detect_encoding(raw)
        return raw.decode(enc, errors="replace")
    except Exception:
        return ""


# ──────────────────────────────────────────────────────────────────────────────
# ВСПОМОГАТЕЛЬНАЯ УТИЛИТА: sniff формата из байт (без Path)
# ──────────────────────────────────────────────────────────────────────────────

def _sniff_suffix_from_bytes(header: bytes, hint_ext: str = "") -> str:
    """
    Определяет суффикс реального формата по первым байтам (без Path).
    Используется внутри zip-рекурсора для вложенных файлов.
    """
    for magic, offset, canon_suffix, _ in _MAGIC_SIGNATURES:
        if header[offset: offset + len(magic)] == magic:
            if canon_suffix == ".zip":
                return _ZIP_OFFICE_MAP.get(hint_ext, (".zip", "ZIP"))[0]
            return canon_suffix
    # HTML-сигнатура в тексте
    try:
        text = header.lstrip(b"\xef\xbb\xbf").decode("utf-8", errors="replace")
        stripped = text.lstrip(" \t\r\n")[:300].lower()
        if any(stripped.startswith(m) for m in ("<!doctype html", "<html", "<head")):
            return ".html"
    except Exception:
        pass
    return hint_ext or ""


# ──────────────────────────────────────────────────────────────────────────────
# НОВЫЕ ЭКСТРАКТОРЫ
# ──────────────────────────────────────────────────────────────────────────────

def extract_strings_binary(path: Path, min_len: int = 6) -> str:
    """
    Извлекает читаемые ASCII-строки из бинарного файла (аналог Unix `strings`).

    Применяется к ELF/PE-файлам и файлам с неизвестным форматом.
    ELF/PE могут содержать строковые константы с ФИО, email, паспортными данными —
    особенно если это утилиты, работающие с персональными данными.

    Фильтрует строки из чистых цифр и повторяющихся символов — они дают
    ложные срабатывания на паттернах вроде паспортов (SNILS, ИНН и т.п.).
    """
    try:
        data = path.read_bytes()[:MAX_TEXT_BYTES]
    except OSError:
        return ""

    # Ищем последовательности печатаемых ASCII-символов длиной >= min_len
    pattern = re.compile(rb'[ -~\t]{' + str(min_len).encode() + rb',}')
    parts = []
    total = 0
    for m in pattern.finditer(data):
        s = m.group(0).decode("ascii", errors="replace").strip()
        # Отсеиваем строки без хотя бы одной буквы — чистые цифровые последовательности
        # дают огромное количество ложных срабатываний паттернов ПДн
        if not any(c.isalpha() for c in s):
            continue
        parts.append(s)
        total += len(s)
        if total >= MAX_TEXT_BYTES:
            break

    return "\n".join(parts)


def extract_xml(path: Path) -> str:
    """
    XML: снимаем теги, возвращаем текстовое содержимое.
    Используем BeautifulSoup если доступен, иначе regex-fallback.
    """
    try:
        raw = path.read_bytes()[:MAX_TEXT_BYTES]
        enc = _detect_encoding(raw)
        text = raw.decode(enc, errors="replace")
    except OSError:
        return ""
    # Пробуем lxml/bs4
    try:
        import bs4
        soup = bs4.BeautifulSoup(text, "lxml-xml")
        return soup.get_text(separator=" ", strip=True)[:MAX_TEXT_BYTES]
    except Exception:
        pass
    # Regex fallback
    return re.sub(r"<[^>]+>", " ", text)[:MAX_TEXT_BYTES]


def extract_gzip(path: Path) -> str:
    """
    GZIP: распаковываем содержимое и сканируем как текст.
    Если внутри другой бинарный формат (например, .tar.gz) — пробуем
    распознать и обработать рекурсивно через временный файл.
    """
    try:
        with gzip.open(path, "rb") as gz:
            raw = gz.read(MAX_TEXT_BYTES)
    except Exception as e:
        log.debug(f"gzip {path.name}: {e}")
        return ""

    inner_suffix = _sniff_suffix_from_bytes(raw[:512])

    # Если внутри ZIP (tar.gz содержит tar, но иногда прямо ZIP)
    if inner_suffix in (".zip", ".docx", ".xlsx"):
        with tempfile.NamedTemporaryFile(suffix=inner_suffix, delete=False) as tmp:
            tmp.write(raw)
            tmp_path = Path(tmp.name)
        try:
            if inner_suffix in (".docx",):
                return extract_docx(tmp_path)
            elif inner_suffix in (".xlsx",):
                return ""  # dict-экстрактор, не текстовый
            else:
                return ""  # ZIP внутри GZ — слишком экзотично
        finally:
            tmp_path.unlink(missing_ok=True)

    # Текст / XML / HTML
    enc = _detect_encoding(raw[:64_000])
    return raw.decode(enc, errors="replace")[:MAX_TEXT_BYTES]


def extract_sqlite(path: Path) -> Dict[str, List[str]]:
    """
    SQLite БД: читаем все таблицы, все строки, все текстовые колонки.
    Поддерживает любую схему — универсальный обход через sqlite3.
    """
    aggregate: Dict[str, List[str]] = defaultdict(list)
    seen: Set[Tuple[str, str]] = set()
    try:
        import sqlite3
        conn = sqlite3.connect(f"file:{path}?mode=ro", uri=True)
        cur = conn.cursor()
        # Получаем список таблиц
        cur.execute("SELECT name FROM sqlite_master WHERE type='table'")
        tables = [row[0] for row in cur.fetchall()]
        for table in tables:
            try:
                cur.execute(f'SELECT * FROM "{table}" LIMIT 50000')
                rows = cur.fetchall()
                buf = []
                for row in rows:
                    buf.append(" | ".join(
                        str(v) for v in row if v is not None
                    ))
                    if len(buf) % 5000 == 0:
                        _merge_findings(aggregate, seen, "\n".join(buf))
                        buf.clear()
                if buf:
                    _merge_findings(aggregate, seen, "\n".join(buf))
            except Exception as e:
                log.debug(f"sqlite table {table}: {e}")
        conn.close()
    except Exception as e:
        log.debug(f"sqlite {path.name}: {e}")
    return dict(aggregate)


# Константы для ZIP-рекурсора
_ZIP_MAX_DEPTH = 3             # максимальная глубина вложения
_ZIP_MAX_MEMBER_BYTES = 50 * 1024 * 1024  # 50 МБ — пропускаем тяжёлые вложения


def extract_zip_recursive(path: Path, _depth: int = 0) -> Dict[str, List[str]]:
    """
    Рекурсивный сканер ZIP-архивов.

    Обходит все вложения, определяет реальный формат каждого по magic bytes,
    применяет подходящий экстрактор. Поддерживает PPTX (XML-слайды), EPUB,
    JAR, вложенные ZIP-архивы. Безопасно: лимит глубины + лимит размера.

    Особо важен для датасетов с подковырками: .pdf содержащий ZIP-архив
    с документами — ни extract_pdf, ни extract_docx его не откроют.
    """
    if _depth > _ZIP_MAX_DEPTH:
        log.debug(f"ZIP рекурсия: превышена глубина {_ZIP_MAX_DEPTH}")
        return {}

    aggregate: Dict[str, List[str]] = defaultdict(list)
    seen: Set[Tuple[str, str]] = set()

    try:
        zf = zipfile.ZipFile(path)
    except zipfile.BadZipFile:
        return {}
    except Exception as e:
        log.debug(f"zip open {path.name}: {e}")
        return {}

    with zf:
        for member in zf.infolist():
            # Пропускаем директории и зашифрованные вложения
            if member.filename.endswith("/"):
                continue
            if member.flag_bits & 0x1:
                log.debug(f"ZIP зашифровано: {member.filename!r}")
                continue
            # Защита от zip-бомб
            if member.file_size > _ZIP_MAX_MEMBER_BYTES:
                log.debug(f"ZIP вложение слишком большое: {member.filename!r} "
                          f"({member.file_size:,} байт)")
                continue

            mem_ext = Path(member.filename).suffix.lower()

            try:
                raw = zf.read(member.filename)
            except Exception as e:
                log.debug(f"ZIP read {member.filename!r}: {e}")
                continue

            real_suffix = _sniff_suffix_from_bytes(raw[:512], mem_ext)

            # Обрабатываем по реальному формату
            try:
                if real_suffix in (".html", ".htm") or mem_ext in (".html", ".htm"):
                    # HTML (например слайды PPTX или страницы EPUB)
                    try:
                        import bs4
                        text = bs4.BeautifulSoup(
                            raw.decode("utf-8", errors="replace"), "html.parser"
                        ).get_text(" ", strip=True)
                    except ImportError:
                        text = re.sub(r"<[^>]+>", " ",
                                      raw.decode("utf-8", errors="replace"))
                    _merge_findings(aggregate, seen, text[:MAX_TEXT_BYTES])

                elif real_suffix == ".xml" or mem_ext == ".xml":
                    text = re.sub(r"<[^>]+>", " ",
                                  raw.decode("utf-8", errors="replace"))
                    _merge_findings(aggregate, seen, text[:MAX_TEXT_BYTES])

                elif real_suffix == ".pdf":
                    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                        tmp.write(raw); tmp_path = Path(tmp.name)
                    try:
                        _merge_findings(aggregate, seen, extract_pdf(tmp_path))
                    finally:
                        tmp_path.unlink(missing_ok=True)

                elif real_suffix == ".docx" or mem_ext == ".docx":
                    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                        tmp.write(raw); tmp_path = Path(tmp.name)
                    try:
                        _merge_findings(aggregate, seen, extract_docx(tmp_path))
                    finally:
                        tmp_path.unlink(missing_ok=True)

                elif real_suffix == ".xlsx" or mem_ext in (".xlsx", ".xls"):
                    with tempfile.NamedTemporaryFile(suffix=mem_ext or ".xlsx",
                                                     delete=False) as tmp:
                        tmp.write(raw); tmp_path = Path(tmp.name)
                    try:
                        sub = extract_xls(tmp_path)
                        for cat, vals in sub.items():
                            for v in vals:
                                k = (cat, v)
                                if k not in seen:
                                    seen.add(k); aggregate[cat].append(v)
                    finally:
                        tmp_path.unlink(missing_ok=True)

                elif real_suffix in (".doc", ".rtf") or mem_ext in (".doc", ".rtf"):
                    with tempfile.NamedTemporaryFile(
                            suffix=real_suffix or mem_ext, delete=False) as tmp:
                        tmp.write(raw); tmp_path = Path(tmp.name)
                    try:
                        _merge_findings(aggregate, seen, extract_doc_rtf(tmp_path))
                    finally:
                        tmp_path.unlink(missing_ok=True)

                elif real_suffix in (".zip", ".docx") and _sniff_suffix_from_bytes(raw[:4]) == ".zip":
                    # Вложенный ZIP (например .jar или .apk)
                    with tempfile.NamedTemporaryFile(suffix=".zip", delete=False) as tmp:
                        tmp.write(raw); tmp_path = Path(tmp.name)
                    try:
                        sub = extract_zip_recursive(tmp_path, _depth + 1)
                        for cat, vals in sub.items():
                            for v in vals:
                                k = (cat, v)
                                if k not in seen:
                                    seen.add(k); aggregate[cat].append(v)
                    finally:
                        tmp_path.unlink(missing_ok=True)

                elif mem_ext in (".csv", ".tsv", ".json", ".jsonl", ".txt", ".md"):
                    enc = _detect_encoding(raw[:64_000])
                    _merge_findings(aggregate, seen,
                                    raw.decode(enc, errors="replace")[:MAX_TEXT_BYTES])

                # Остальное (изображения, бинарники) — пропускаем
                # (для OCR нужен Tesseract + PIL, что требует файла на диске;
                # пропуск оправдан — текст в архиве важнее TIFF-сканов)

            except Exception as e:
                log.debug(f"ZIP member {member.filename!r}: {e}")

    return dict(aggregate)


def extract_appended_data(path: Path, real_format: str) -> str:
    """
    Извлекает данные после маркера конца файла (DATA_AFTER_EOF).

    JPEG: после FFD9 (End Of Image)
    PDF:  после последнего %%EOF
    PNG:  после IEND chunk

    Appended-данные — классический способ спрятать текстовый документ
    (или другой файл) внутри изображения. Если хвост содержит текст —
    он будет пропущен при обычном OCR/извлечении.
    """
    try:
        data = path.read_bytes()
    except OSError:
        return ""

    tail = b""

    if real_format in ("JPEG", ".jpg"):
        eoi = data.rfind(b"\xff\xd9")
        if eoi != -1 and eoi + 2 < len(data):
            tail = data[eoi + 2:]

    elif real_format in ("PDF", ".pdf"):
        eof_pos = data.rfind(b"%%EOF")
        if eof_pos != -1:
            raw_tail = data[eof_pos + 5:].lstrip(b"\r\n \x00")
            if raw_tail:
                tail = raw_tail

    elif real_format in ("PNG", ".png"):
        # PNG заканчивается chunk IEND: 00 00 00 00 49 45 4E 44 AE 42 60 82
        iend = data.rfind(b"IEND\xaeB`\x82")
        if iend != -1 and iend + 8 < len(data):
            tail = data[iend + 8:]

    if not tail or len(tail) < 8:
        return ""

    # Определяем что в хвосте: текст или вложенный файл
    real_tail_suffix = _sniff_suffix_from_bytes(tail[:512])

    if real_tail_suffix in (".html", ".htm"):
        try:
            import bs4
            return bs4.BeautifulSoup(
                tail.decode("utf-8", errors="replace"), "html.parser"
            ).get_text(" ", strip=True)[:MAX_TEXT_BYTES]
        except Exception:
            pass

    if real_tail_suffix == ".pdf":
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(tail); tmp_path = Path(tmp.name)
        try:
            return extract_pdf(tmp_path)
        finally:
            tmp_path.unlink(missing_ok=True)

    if real_tail_suffix in (".zip", ".docx"):
        with tempfile.NamedTemporaryFile(suffix=".zip", delete=False) as tmp:
            tmp.write(tail); tmp_path = Path(tmp.name)
        try:
            sub = extract_zip_recursive(tmp_path)
            # Конвертируем dict в плоский текст для слияния с основным результатом
            parts = []
            for vals in sub.values():
                parts.extend(vals)
            return "\n".join(parts)
        finally:
            tmp_path.unlink(missing_ok=True)

    # Пробуем декодировать как текст
    enc = _detect_encoding(tail[:64_000])
    try:
        decoded = tail.decode(enc, errors="replace")
        # Возвращаем только если содержит читаемый текст (не бинарный мусор)
        printable_ratio = sum(1 for c in decoded[:1000] if c.isprintable()) / min(1000, len(decoded))
        if printable_ratio > 0.7:
            return decoded[:MAX_TEXT_BYTES]
    except Exception:
        pass

    return ""


# ──────────────────────────────────────────────────────────────────────────────
# УЛУЧШЕНИЕ EXTRACT_PDF: вложения + JS-содержимое
# ──────────────────────────────────────────────────────────────────────────────

def _extract_pdf_attachments(path: Path) -> str:
    """
    Извлекает и сканирует файловые вложения из PDF (/EmbeddedFiles).
    Вызывается из extract_pdf после основного текста.
    Требует PyMuPDF (fitz).
    """
    parts = []
    try:
        import fitz
        with fitz.open(str(path)) as doc:
            if doc.embfile_count() == 0:
                return ""
            for i in range(doc.embfile_count()):
                info = doc.embfile_info(i)
                name = info.get("filename", f"attachment_{i}")
                att_ext = Path(name).suffix.lower()
                try:
                    raw = doc.embfile_get(i)
                    if not raw:
                        continue
                except Exception:
                    continue

                real_suffix = _sniff_suffix_from_bytes(raw[:512], att_ext)

                # Обрабатываем по реальному формату вложения
                with tempfile.NamedTemporaryFile(
                        suffix=real_suffix or att_ext or ".bin", delete=False) as tmp:
                    tmp.write(raw); tmp_path = Path(tmp.name)
                try:
                    if real_suffix == ".pdf" or att_ext == ".pdf":
                        text = extract_pdf(tmp_path)
                    elif real_suffix in (".docx",) or att_ext == ".docx":
                        text = extract_docx(tmp_path)
                    elif real_suffix in (".doc", ".rtf") or att_ext in (".doc", ".rtf"):
                        text = extract_doc_rtf(tmp_path)
                    elif real_suffix in (".xlsx", ".xls") or att_ext in (".xlsx", ".xls"):
                        sub = extract_xls(tmp_path)
                        text = " ".join(
                            v for vals in sub.values() for v in vals
                        )
                    elif real_suffix in (".html", ".htm") or att_ext in (".html", ".htm"):
                        text = extract_html(tmp_path)
                    elif real_suffix in (".zip",) or att_ext == ".zip":
                        sub = extract_zip_recursive(tmp_path)
                        text = " ".join(v for vals in sub.values() for v in vals)
                    else:
                        enc = _detect_encoding(raw[:64_000])
                        text = raw.decode(enc, errors="replace")
                    if text:
                        parts.append(f"[вложение: {name}]\n{text}")
                except Exception as e:
                    log.debug(f"PDF вложение {name!r}: {e}")
                finally:
                    tmp_path.unlink(missing_ok=True)
    except ImportError:
        pass
    except Exception as e:
        log.debug(f"PDF attachments {path.name}: {e}")
    return "\n".join(parts)[:MAX_TEXT_BYTES]


# ──────────────────────────────────────────────────────────────────────────────
# УЛУЧШЕНИЕ EXTRACT_DOCX: OLE-embedded объекты
# ──────────────────────────────────────────────────────────────────────────────

def _extract_docx_embedded(path: Path) -> str:
    """
    Извлекает OLE-embedded объекты из DOCX.

    DOCX — это ZIP. Внутри могут быть word/embeddings/Microsoft_Word_Document*.bin
    (или *.docx, *.xlsx) — это OLE-вставки из других документов.
    python-docx их игнорирует. Мы сканируем вручную.
    """
    parts = []
    try:
        zf = zipfile.ZipFile(path)
    except Exception:
        return ""
    with zf:
        for member in zf.infolist():
            name = member.filename
            # Ищем папку embeddings — там лежат вставленные объекты
            if "embeddings" not in name.lower():
                continue
            if member.filename.endswith("/"):
                continue
            try:
                raw = zf.read(name)
            except Exception:
                continue

            att_ext = Path(name).suffix.lower()
            real_suffix = _sniff_suffix_from_bytes(raw[:512], att_ext)

            with tempfile.NamedTemporaryFile(
                    suffix=real_suffix or att_ext or ".bin", delete=False) as tmp:
                tmp.write(raw); tmp_path = Path(tmp.name)
            try:
                if real_suffix in (".doc", ".rtf") or att_ext in (".doc", ".bin"):
                    text = extract_doc_rtf(tmp_path)
                elif real_suffix == ".pdf":
                    text = extract_pdf(tmp_path)
                elif real_suffix in (".xlsx", ".xls"):
                    sub = extract_xls(tmp_path)
                    text = " ".join(v for vals in sub.values() for v in vals)
                else:
                    enc = _detect_encoding(raw[:64_000])
                    text = raw.decode(enc, errors="replace")
                if text.strip():
                    parts.append(f"[embedded: {Path(name).name}]\n{text}")
            except Exception as e:
                log.debug(f"DOCX embedded {name!r}: {e}")
            finally:
                tmp_path.unlink(missing_ok=True)

    return "\n".join(parts)[:MAX_TEXT_BYTES]


# Специальный маркер для случая «OCR недоступен».
# Возвращается экстрактором, а analyze_file ставит особую рекомендацию.
OCR_UNAVAILABLE_SENTINEL = "__OCR_NOT_AVAILABLE__"


def _prepare_ocr_variants_pil(img):
    """Готовит варианты изображения для Tesseract (CPU-бэкенд).

    EasyOCR на GPU сам делает препроцессинг внутри своего детектора, ему
    варианты не нужны — это было бы пустой тратой VRAM.
    """
    from PIL import ImageOps, ImageFilter

    img = ImageOps.exif_transpose(img)
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")

    gray = ImageOps.grayscale(img)

    # Для маленьких скриншотов/кадров поднимаем разрешение — Tesseract заметно лучше
    # распознаёт текст на документах после upscale.
    max_side = max(gray.size)
    scale = 1
    if max_side < 900:
        scale = 4
    elif max_side < 1400:
        scale = 3
    elif max_side < 2000:
        scale = 2
    if scale > 1:
        gray = gray.resize((gray.width * scale, gray.height * scale))

    base = ImageOps.autocontrast(gray)
    sharp = base.filter(ImageFilter.UnsharpMask(radius=1.5, percent=180, threshold=3))
    bw = sharp.point(lambda p: 255 if p >= 170 else 0, mode="1").convert("L")

    variants = []
    seen = set()
    for candidate in (base, sharp, bw):
        key = (candidate.size, candidate.tobytes())
        if key in seen:
            continue
        seen.add(key)
        variants.append(candidate)
    return variants


def _ocr_pil_image_easyocr(img, reader) -> str:
    """Быстрый путь: EasyOCR на GPU. Один вызов — и детекция, и распознавание.

    Внутри EasyOCR сам батчит нарезанные текстовые регионы (см. параметр
    batch_size). Ему лучше всего скормить картинку покрупнее (до ~2000 px
    по большей стороне) — детектор в upscale не нуждается.
    """
    try:
        import numpy as np  # type: ignore
        from PIL import ImageOps
    except ImportError:
        return ""

    try:
        img = ImageOps.exif_transpose(img)
        if img.mode != "RGB":
            img = img.convert("RGB")

        # Маленькие скриншоты всё-таки поднимаем — иначе детектор пропускает
        # мелкий шрифт (характерно для кадров из MP4 640×360).
        max_side = max(img.size)
        if max_side < 800:
            scale = 2
            img = img.resize((img.width * scale, img.height * scale))

        arr = np.array(img)
        # detail=1 → возвращает [bbox, text, conf] для каждого региона
        # paragraph=False — paragraph=True часто теряет регионы на сложных сканах
        # Фильтруем по confidence >= 0.2 чтобы отсеять явный мусор
        results = reader.readtext(
            arr,
            detail=1,
            paragraph=False,
            batch_size=_GPU_BATCH_SIZE,
        )
        lines = [r[1] for r in results if r[2] >= 0.2 and r[1].strip()]
        text = ("\n".join(lines))[:MAX_TEXT_BYTES]

        # Освобождаем VRAM после каждого файла — EasyOCR не делает это сам
        try:
            import torch  # type: ignore
            if torch.cuda.is_available():
                torch.cuda.empty_cache()
        except Exception:
            pass

        return text
    except Exception as e:
        log.debug(f"easyocr: {e}")
        try:
            import torch  # type: ignore
            if torch.cuda.is_available():
                torch.cuda.empty_cache()
        except Exception:
            pass
        return ""


def _ocr_pil_image_paddleocr(img, ocr) -> str:
    """PaddleOCR PP-OCRv5 на GPU — обычно в 3-10× быстрее EasyOCR.

    Поддерживает оба распространённых API:
      • v3.x: ocr.predict(arr) → список объектов с атрибутом rec_texts
      • v2.x: ocr.ocr(arr, cls=True) → [[[bbox], (text, conf)], ...]
    """
    try:
        import numpy as np  # type: ignore
        from PIL import ImageOps
    except ImportError:
        return ""

    try:
        img = ImageOps.exif_transpose(img)
        if img.mode != "RGB":
            img = img.convert("RGB")

        arr = np.array(img)

        # === API 3.x (новый) ===
        predict_fn = getattr(ocr, "predict", None)
        if predict_fn is not None:
            try:
                result = predict_fn(arr)
                texts: List[str] = []
                if result:
                    items = result if isinstance(result, list) else [result]
                    for item in items:
                        # Вариант 1: атрибут rec_texts (mobile API)
                        rec_texts = getattr(item, "rec_texts", None)
                        if rec_texts:
                            texts.extend([t for t in rec_texts if t])
                            continue
                        # Вариант 2: dict с rec_texts
                        if isinstance(item, dict):
                            rt = item.get("rec_texts") or item.get("rec_text")
                            if rt:
                                texts.extend(rt if isinstance(rt, list) else [str(rt)])
                                continue
                        # Вариант 3: server модель возвращает список
                        # словарей [{rec_text, rec_score, ...}, ...]
                        if isinstance(item, list):
                            for sub in item:
                                if isinstance(sub, dict):
                                    t = sub.get("rec_text") or sub.get("text")
                                    if t:
                                        texts.append(str(t))
                        # Вариант 4: итерируемый объект с rec_text
                        try:
                            for sub in item:
                                t = getattr(sub, "rec_text", None)
                                if t:
                                    texts.append(str(t))
                        except TypeError:
                            pass
                # Диагностика: логируем структуру если текст не извлечён
                if not texts and result:
                    items = result if isinstance(result, list) else [result]
                    first = items[0] if items else None
                    log.debug(
                        f"paddleocr predict: текст не извлечён. "
                        f"type={type(first).__name__}, "
                        f"attrs={[a for a in dir(first) if not a.startswith('_')][:8] if first else []}"
                    )
                if texts:
                    return ("\n".join(texts))[:MAX_TEXT_BYTES]
            except Exception as e:
                log.debug(f"paddleocr predict: {e}")

        # === API 2.x (старый) ===
        try:
            result = ocr.ocr(arr, cls=True)
        except TypeError:
            result = ocr.ocr(arr)

        texts = []
        # result может быть [[...]] (один элемент внешнего списка на изображение)
        # или сразу [...]. Нормализуем.
        if result and isinstance(result, list):
            page_results = result[0] if (len(result) == 1 and
                                         isinstance(result[0], list)) else result
            for line in (page_results or []):
                # line = [ [bbox_points], (text, confidence) ]
                if isinstance(line, (list, tuple)) and len(line) >= 2:
                    rec = line[1]
                    if isinstance(rec, (list, tuple)) and len(rec) >= 1:
                        text = rec[0]
                        if isinstance(text, str):
                            texts.append(text)
        return ("\n".join(texts))[:MAX_TEXT_BYTES]
    except Exception as e:
        log.debug(f"paddleocr: {e}")
        return ""


def _ocr_pil_image_surya(img, state: dict) -> str:
    """Surya OCR — PyTorch-based, работает на GPU без проблем с CUDNN.

    Surya автоматически определяет язык (90+ языков включая ru/en/de/cs/fr).
    Скорость на GPU ~10-20× быстрее Tesseract.
    """
    try:
        from PIL import Image as PILImage  # type: ignore
        det = state["det"]
        rec = state["rec"]

        if not isinstance(img, PILImage.Image):
            return ""
        if img.mode != "RGB":
            img = img.convert("RGB")

        # Surya: детекция текстовых строк
        predictions = det([img])
        if not predictions or not predictions[0].bboxes:
            return ""

        # Распознавание — langs=None означает автодетект языка
        rec_predictions = rec([img], [None], det_predictions=predictions)
        if not rec_predictions:
            return ""

        texts = []
        for page in rec_predictions:
            for line in page.text_lines:
                if line.text and line.text.strip():
                    texts.append(line.text.strip())
        return ("\n".join(texts))[:MAX_TEXT_BYTES]
    except Exception as e:
        log.debug(f"surya ocr: {e}")
        return ""


def _ocr_pil_image(img, lang: Optional[str] = None) -> str:
    """Главная точка входа для OCR одной картинки. Выбирает бэкенд по глобалу
    _OCR_ENGINE, который инициализируется в _init_worker().
    """
    if _OCR_ENGINE is None:
        _init_ocr_engine("auto")

    engine, state = _OCR_ENGINE  # type: ignore[misc]
    if engine == "none":
        return OCR_UNAVAILABLE_SENTINEL
    if engine == "surya":
        return _ocr_pil_image_surya(img, state)
    if engine == "paddleocr":
        return _ocr_pil_image_paddleocr(img, state)
    if engine == "easyocr":
        return _ocr_pil_image_easyocr(img, state)

    # === Tesseract (CPU-fallback): старый путь с несколькими вариантами ===
    try:
        import pytesseract
    except ImportError:
        return OCR_UNAVAILABLE_SENTINEL

    effective_lang = lang or _OCR_LANG_TESS
    texts = []
    seen = set()
    configs = ("--oem 1 --psm 6", "--oem 1 --psm 11")

    for variant in _prepare_ocr_variants_pil(img):
        for cfg in configs:
            try:
                t = pytesseract.image_to_string(
                    variant, lang=effective_lang, config=cfg, timeout=20
                )
            except pytesseract.TesseractNotFoundError:
                return OCR_UNAVAILABLE_SENTINEL
            except RuntimeError:
                continue
            except Exception:
                continue

            norm = re.sub(r"\s+", " ", t).strip()
            if not norm:
                continue
            if norm in seen:
                continue
            seen.add(norm)
            texts.append(t)

    return "\n".join(texts)[:MAX_TEXT_BYTES]


def _preprocess_image_for_ocr(img) -> "Image.Image":
    """Унифицированная предобработка изображения перед OCR.

    Стратегия (твоя идея):
      1. EXIF-ротация
      2. Конвертация в RGB (TIF/BMP/GIF могут быть Palette/CMYK/RGBA)
      3. Downscale если > MAX_OCR_PX по длинной стороне — OCR не нужно
         более 2000-3000 px, а огромные TIF (~5700px) замедляют в разы
      4. Upscale если < MIN_OCR_PX — мелкий скан читается хуже
    """
    from PIL import Image, ImageOps
    MAX_OCR_PX = 2500  # px по длинной стороне — оптимум для PaddleOCR
    MIN_OCR_PX = 800

    img = ImageOps.exif_transpose(img)
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")

    w, h = img.size
    long_side = max(w, h)

    if long_side > MAX_OCR_PX:
        scale = MAX_OCR_PX / long_side
        img = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
    elif long_side < MIN_OCR_PX:
        scale = MIN_OCR_PX / long_side
        img = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)

    return img


def extract_image_tesseract(path: Path) -> str:
    """OCR изображения (TIF/JPEG/PNG/GIF/BMP).

    Стратегия предобработки (ускорение ~3-5x на TIF):
      - TIF/TIFF конвертируется в RGB с downscale до 2500px (было 5700+px)
      - Многостраничные TIFF/GIF читаются пофреймово (макс. 10 страниц)
      - Все форматы проходят через _preprocess_image_for_ocr перед OCR
    """
    try:
        from PIL import Image, ImageSequence
    except ImportError:
        return OCR_UNAVAILABLE_SENTINEL

    if _OCR_ENGINE is None:
        _init_ocr_engine("auto")
    if _ocr_engine_name() == "none":
        return OCR_UNAVAILABLE_SENTINEL

    try:
        img = Image.open(path)
        texts = []
        max_frames = 10

        if getattr(img, "n_frames", 1) > 1:
            frames = list(ImageSequence.Iterator(img))[:max_frames]
        else:
            frames = [img]

        for frame in frames:
            # Ключевая оптимизация: конвертируем/сжимаем ДО OCR
            frame = frame.copy()  # avoid iterator invalidation
            frame = _preprocess_image_for_ocr(frame)
            t = _ocr_pil_image(frame)
            if t == OCR_UNAVAILABLE_SENTINEL:
                return OCR_UNAVAILABLE_SENTINEL
            if t and t.strip():
                texts.append(t)

        return "\n".join(texts)[:MAX_TEXT_BYTES]
    except Exception as e:
        log.debug(f"ocr image {path.name}: {e}")
        return ""


VIDEO_UNAVAILABLE_SENTINEL = "__VIDEO_NOT_AVAILABLE__"


def cv2_available() -> bool:
    """Проверяет доступность OpenCV (cv2) для обработки видео."""
    try:
        import cv2  # noqa: F401
        return True
    except ImportError:
        return False


def extract_mp4_frames_ocr(path: Path) -> str:
    """Видео (MP4/AVI/MOV): извлекаем кадры чаще и OCR'им так же, как фото.

    В исходной версии было две проблемы:
      1) редкая выборка кадров (кадр с документом легко пропускался);
      2) слабый препроцессинг — размытые скриншоты и записи экрана почти не читались.
    """
    try:
        import cv2
        import numpy as np
        from PIL import Image
    except ImportError:
        return VIDEO_UNAVAILABLE_SENTINEL

    # OCR-бэкенд (EasyOCR/Tesseract) инициализируется в _init_worker().
    # Здесь просто проверяем, что он вообще доступен.
    if _OCR_ENGINE is None:
        _init_ocr_engine("auto")
    if _ocr_engine_name() == "none":
        return OCR_UNAVAILABLE_SENTINEL

    MAX_VIDEO_FRAMES = 12   # документы статичны — 12 кадров достаточно
    MIN_INTERVAL_SEC = 2.0  # кадр каждые 2 сек минимум
    MAX_INTERVAL_SEC = 5.0

    try:
        cap = cv2.VideoCapture(str(path))
        if not cap.isOpened():
            return ""

        total = int(cap.get(cv2.CAP_PROP_FRAME_COUNT) or 0)
        fps = cap.get(cv2.CAP_PROP_FPS) or 25.0
        dur_s = total / fps if fps > 0 else 0

        if total <= 0 or dur_s <= 0:
            cap.release()
            return ""

        interval_s = min(MAX_INTERVAL_SEC, max(MIN_INTERVAL_SEC, dur_s / MAX_VIDEO_FRAMES))
        interval_f = max(1, int(round(interval_s * fps)))

        frame_indices = set(range(0, total, interval_f))
        frame_indices.update({0, max(0, total // 2), max(0, total - 1)})
        frame_indices = sorted(frame_indices)[:MAX_VIDEO_FRAMES]

        texts = []
        prev_fp = None

        for idx in frame_indices:
            cap.set(cv2.CAP_PROP_POS_FRAMES, float(idx))
            ok, frame = cap.read()
            if not ok or frame is None:
                continue

            small = cv2.resize(frame, (16, 16))
            gray_small = cv2.cvtColor(small, cv2.COLOR_BGR2GRAY)
            fp = gray_small > gray_small.mean()
            if prev_fp is not None:
                dist = int(np.count_nonzero(fp != prev_fp))
                if dist <= 6:
                    continue
            prev_fp = fp

            img = Image.fromarray(cv2.cvtColor(frame, cv2.COLOR_BGR2RGB))
            img = _preprocess_image_for_ocr(img)  # downscale до 2500px
            t = _ocr_pil_image(img)
            if t == OCR_UNAVAILABLE_SENTINEL:
                cap.release()
                return OCR_UNAVAILABLE_SENTINEL
            if t and t.strip():
                texts.append(t)

        cap.release()
        return "\n".join(texts)[:MAX_TEXT_BYTES]

    except Exception as e:
        log.debug(f"mp4 {path.name}: {e}")
        return ""


# Сопоставление расширений → (формат, «тип обработки»)
# тип: "text" (можно в ProcessPool), "struct" (тоже), "ocr" (отдельный пул)
EXT_HANDLERS: Dict[str, Tuple[str, str]] = {
    ".csv":     ("CSV",     "struct"),
    ".tsv":     ("TSV",     "struct"),
    ".json":    ("JSON",    "struct"),
    ".jsonl":   ("JSONL",   "struct"),
    ".parquet": ("Parquet", "struct"),
    ".xls":     ("XLS",     "struct"),
    ".xlsx":    ("XLSX",    "struct"),
    ".zip":     ("ZIP",     "struct"),   # архив — сканируем рекурсивно
    ".sqlite":  ("SQLite",  "struct"),
    ".db":      ("SQLite",  "struct"),
    ".pdf":     ("PDF",     "text"),
    ".doc":     ("DOC",     "text"),
    ".docx":    ("DOCX",    "text"),
    ".rtf":     ("RTF",     "text"),
    ".html":    ("HTML",    "text"),
    ".htm":     ("HTML",    "text"),
    ".xml":     ("XML",     "text"),    # XML — как текст со strip тегов
    ".txt":     ("TXT",     "text"),
    ".md":      ("MD",      "text"),
    ".gz":      ("GZIP",    "text"),    # gzip — распаковываем и сканируем
    ".elf":     ("ELF",     "text"),    # бинарники — strings-извлечение
    ".exe":     ("PE",      "text"),
    ".tif":     ("TIF",     "ocr"),
    ".tiff":    ("TIFF",    "ocr"),
    ".jpg":     ("JPEG",    "ocr"),
    ".jpeg":    ("JPEG",    "ocr"),
    ".png":     ("PNG",     "ocr"),
    ".gif":     ("GIF",     "ocr"),
    ".bmp":     ("BMP",     "ocr"),
    ".mp4":     ("MP4",     "ocr"),
    ".avi":     ("AVI",     "ocr"),
    ".mov":     ("MOV",     "ocr"),
}


# ──────────────────────────────────────────────────────────────────────────────
# ОПРЕДЕЛЕНИЕ РЕАЛЬНОГО ФОРМАТА ПО СОДЕРЖИМОМУ (MAGIC BYTES)
# ──────────────────────────────────────────────────────────────────────────────
#
# Задание с подковыркой: некоторые файлы в датасете имеют «неправильное»
# расширение. Например, файл с расширением .pdf внутри является HTML, а .docx
# может оказаться OLE2-документом (старый Word .doc), HTML или даже изображением.
# Анализировать такие файлы «наивным» способом (только по суффиксу) — значит
# либо ничего не найти (PyMuPDF не откроет HTML как PDF), либо пропустить ПДн.
#
# Стратегия:
#   1. Читаем первые 512 байт файла.
#   2. Проверяем точные сигнатуры (magic bytes).
#   3. Если сигнатура не совпала — ищем HTML-маркеры в тексте заголовка.
#   4. Если реальный формат отличается от расширения → логируем предупреждение
#      и перегружаем маршрутизацию.
#
# Поддерживаемые сигнатуры:
#   %PDF       → PDF
#   PK\x03\x04 → ZIP-контейнер (DOCX / XLSX / PPTX / ODT / …)
#   D0 CF 11 E0 → OLE2 (старый DOC / XLS / PPT)
#   \x89PNG    → PNG
#   FF D8 FF   → JPEG
#   GIF8[79]a  → GIF
#   BM         → BMP
#   II*\x00 / MM\x00* → TIFF
#   {\rtf      → RTF
#   <?xml      → XML
#   <!DOCTYPE html / <html → HTML  (также ищем в первых 512 байтах)

# (magic_bytes, offset_from_start, canonical_suffix, format_label)
_MAGIC_SIGNATURES: List[Tuple[bytes, int, str, str]] = [
    (b"%PDF",                             0, ".pdf",    "PDF"),
    (b"PK\x03\x04",                      0, ".zip",    "ZIP"),   # ZIP-контейнер — уточним ниже
    (b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1", 0, ".doc",  "OLE2"),  # DOC/XLS/PPT (binary)
    (b"\x89PNG\r\n\x1a\n",               0, ".png",    "PNG"),
    (b"\xff\xd8\xff",                    0, ".jpg",    "JPEG"),
    (b"GIF87a",                          0, ".gif",    "GIF"),
    (b"GIF89a",                          0, ".gif",    "GIF"),
    (b"BM",                              0, ".bmp",    "BMP"),
    (b"II\x2a\x00",                      0, ".tif",    "TIFF"),
    (b"MM\x00\x2a",                      0, ".tif",    "TIFF"),
    (b"{\\rtf",                          0, ".rtf",    "RTF"),
    (b"<?xml",                           0, ".xml",    "XML"),
    (b"\x1f\x8b",                        0, ".gz",     "GZIP"),   # gzip
    (b"7z\xbc\xaf'\x1c",               0, ".7z",     "7ZIP"),   # 7-Zip
    (b"Rar!\x1a\x07",                   0, ".rar",    "RAR"),    # RAR
    (b"SQLite format 3\x00",             0, ".sqlite", "SQLite3"),# SQLite
    (b"\x7fELF",                         0, ".elf",    "ELF"),    # Linux ELF
    (b"MZ",                              0, ".exe",    "PE"),     # Windows PE/EXE
]

# Расширения, которые являются ZIP-контейнерами (OOXML/ODF).
# Используем расширение файла чтобы отличить DOCX от XLSX внутри ZIP.
_ZIP_OFFICE_MAP: Dict[str, Tuple[str, str]] = {
    ".docx": (".docx", "DOCX"),
    ".odt":  (".docx", "DOCX"),   # обрабатываем через python-docx где возможно
    ".xlsx": (".xlsx", "XLSX"),
    ".ods":  (".xlsx", "XLSX"),
    ".pptx": (".zip",  "ZIP"),    # → рекурсивный ZIP-сканер извлечёт XML-слайды
    ".odp":  (".zip",  "ZIP"),
    ".zip":  (".zip",  "ZIP"),    # явный .zip → рекурсивный сканер
    ".jar":  (".zip",  "ZIP"),
    ".epub": (".zip",  "ZIP"),
}


def sniff_real_format(path: Path) -> Tuple[str, str]:
    """
    Определяет реальный формат файла по магическим байтам и HTML-сигнатурам.

    Возвращает ``(suffix, format_label)`` — в том же формате, что EXT_HANDLERS.
    Если реальный формат определить не удалось, возвращает суффикс расширения.

    Алгоритм:
      1. Читаем первые 512 байт.
      2. Проверяем точные magic bytes из _MAGIC_SIGNATURES.
      3. Для ZIP — уточняем тип по расширению файла (DOCX / XLSX / ...).
      4. Для OLE2 — всегда возвращаем ".doc" независимо от расширения.
      5. Если magic bytes не совпали — пробуем декодировать заголовок как текст
         и ищем HTML-маркеры: <!doctype html, <html, <head, <body.
         (HTML-экспорт из Word/Excel часто содержит эти теги.)
      6. Fallback: доверяем расширению файла.
    """
    ext = path.suffix.lower()
    fallback_fmt = EXT_HANDLERS.get(ext, ("UNKNOWN", "skip"))[0]

    try:
        header = path.read_bytes()[:512]
    except OSError:
        return ext, fallback_fmt

    # 1. Magic bytes
    for magic, offset, canon_suffix, label in _MAGIC_SIGNATURES:
        if header[offset: offset + len(magic)] == magic:

            # ZIP-контейнер → уточняем по расширению
            if canon_suffix == ".zip":
                if ext in _ZIP_OFFICE_MAP:
                    return _ZIP_OFFICE_MAP[ext]
                # Неизвестный ZIP-файл → рекурсивный ZIP-сканер
                return ".zip", "ZIP"

            return canon_suffix, label

    # 2. HTML-сигнатуры в заголовке (без BOM, без ведущих пробелов)
    try:
        # Убираем UTF-8 BOM если есть
        text_head = header.lstrip(b"\xef\xbb\xbf").decode("utf-8", errors="replace")
    except Exception:
        text_head = ""

    stripped = text_head.lstrip(" \t\r\n")
    lower512 = stripped[:512].lower()

    html_markers = ("<!doctype html", "<html", "<head", "<body")
    if any(lower512.startswith(m) for m in html_markers) or "<html" in lower512:
        return ".html", "HTML"

    # 3. Fallback — доверяем расширению
    return ext, fallback_fmt


# ──────────────────────────────────────────────────────────────────────────────
# АНАЛИЗ ОДНОГО ФАЙЛА
# ──────────────────────────────────────────────────────────────────────────────

def _merge_findings(aggregate: Dict[str, List[str]],
                    seen: Set[Tuple[str, str]],
                    text: str):
    """Ищет ПДн в тексте и сливает в общий агрегат, с дедупликацией."""
    if not text:
        return
    found = detect_pdn(text)
    for cat, values in found.items():
        for v in values:
            key = (cat, v)
            if key in seen:
                continue
            seen.add(key)
            aggregate[cat].append(v)


def _merge_findings_dict_inline(base: Dict[str, List[str]],
                                extra: Dict[str, List[str]]):
    """
    Сливает extra в base (оба — результаты detect_pdn).
    Используется для объединения основного результата с результатом
    хвостового сканирования (DATA_AFTER_EOF) или вложений.
    Дедупликация по (категория, значение).
    """
    for cat, vals in extra.items():
        existing = set(base.get(cat, []))
        for v in vals:
            if v not in existing:
                base.setdefault(cat, []).append(v)
                existing.add(v)


RECOMMENDATIONS = {
    "УЗ-1": "Спец. категории / биометрия. Требуются меры 1-го уровня защищённости "
            "(Приказ ФСТЭК № 21). Шифрование, аттестация ИС, строгий аудит.",
    "УЗ-2": "Платёжные данные или массив гос. идентификаторов. Меры 2-го уровня: "
            "контроль доступа, шифрование, учёт носителей. PCI DSS при карточных данных.",
    "УЗ-3": "Гос. идентификаторы / массив обычных ПДн. Меры 3-го уровня: ролевая "
            "модель, бэкапы, обучение персонала, инвентаризация.",
    "УЗ-4": "Только обычные ПДн в небольшом объёме. Базовые меры: разграничение "
            "прав, антивирус, обновления ПО, согласия субъектов.",
    "—":   "ПДн не обнаружены. Дополнительных мер не требуется.",
}


@dataclass
class FileResult:
    path: str
    format: str
    size_bytes: int
    pdn_categories: Dict[str, List[str]] = field(default_factory=dict)
    total_findings: int = 0
    uz_level: str = "—"
    recommendation: str = ""
    error: Optional[str] = None
    elapsed_sec: float = 0.0
    # Заполняется если реальный формат файла отличается от расширения.
    # Например: расширение «.pdf» → реальный формат «HTML».
    real_format: Optional[str] = None
    # Флаг для двухфазного OCR: PDF без text-layer, прошедший через CPU-пул
    # (фаза 1) помечается pending_ocr=True и переотправляется в GPU OCR-пул
    # (фаза 2). Нельзя пытаться поднять EasyOCR в 16 CPU-воркерах — VRAM OOM.
    pending_ocr: bool = False


def analyze_file(path_str: str, allow_ocr: bool = True) -> Dict[str, Any]:
    """
    Анализирует один файл. Возвращает dict (сериализуется в JSON).
    Выполняется в ProcessPoolExecutor → не использовать global state.
    """
    t0 = time.monotonic()
    path = Path(path_str)
    ext = path.suffix.lower()          # расширение из имени файла
    fmt, kind = EXT_HANDLERS.get(ext, ("UNKNOWN", "skip"))

    try:
        size = path.stat().st_size
    except OSError as e:
        return asdict(FileResult(
            path=str(path), format=fmt, size_bytes=0,
            error=f"stat failed: {e}",
            elapsed_sec=time.monotonic() - t0,
        ))

    # ── Верификация реального формата по содержимому ──────────────────────────
    # Задание с подковыркой: некоторые файлы имеют «чужое» расширение.
    # sniff_real_format() читает magic bytes и HTML-сигнатуры и возвращает
    # реальный суффикс. Если он отличается от расширения — перегружаем fmt/kind.
    real_suffix, real_fmt_label = sniff_real_format(path)
    real_format_note: Optional[str] = None          # запишем в отчёт если подмена

    if real_suffix != ext:
        real_format_note = (
            f"{real_fmt_label} (расширение: {ext!r} → реальный: {real_suffix!r})"
        )
        log.debug(
            f"Подмена формата: {path.name!r}  "            f"расширение={ext!r}  реальный={real_suffix!r}"
        )
        # Переключаемся на реальный формат для всей дальнейшей обработки
        suffix = real_suffix
        fmt    = real_fmt_label + f" [заявлен как {ext!r}]"
        kind   = EXT_HANDLERS.get(real_suffix, ("UNKNOWN", "skip"))[1]
    else:
        suffix = ext
    # ─────────────────────────────────────────────────────────────────────────

    # Пропуск тяжёлых изображений
    if kind == "ocr" and size > MAX_IMAGE_BYTES and suffix in (
        ".tif", ".tiff", ".jpg", ".jpeg", ".png", ".bmp"
    ):
        return asdict(FileResult(
            path=str(path), format=fmt, size_bytes=size,
            recommendation=f"Пропущено: изображение > {MAX_IMAGE_BYTES // (1024*1024)} МБ",
            real_format=real_format_note,
            elapsed_sec=time.monotonic() - t0,
        ))

    if kind == "ocr" and not allow_ocr:
        return asdict(FileResult(
            path=str(path), format=fmt, size_bytes=size,
            recommendation="OCR отключён (--no-ocr)",
            real_format=real_format_note,
            elapsed_sec=time.monotonic() - t0,
        ))

    # Диспатч
    ocr_unavailable = False
    try:
        if suffix == ".csv" or suffix == ".tsv":
            found = extract_csv(path)
            text_for_filter = ""
        elif suffix in (".json", ".jsonl"):
            found = extract_json(path)
            text_for_filter = ""
        elif suffix == ".parquet":
            found = extract_parquet(path)
            text_for_filter = ""
        elif suffix in (".xls", ".xlsx"):
            found = extract_xls(path)
            text_for_filter = ""
        elif suffix in (".sqlite", ".db"):
            # SQLite БД — читаем все текстовые колонки
            found = extract_sqlite(path)
            text_for_filter = ""
        elif suffix == ".zip":
            # ZIP-архив (включая PPTX, EPUB, JAR и прочие ZIP-контейнеры)
            # — рекурсивный обход всех вложений
            found = extract_zip_recursive(path)
            text_for_filter = ""
        elif suffix == ".pdf":
            text_for_filter = extract_pdf(path)   # уже включает вложения

            # PDF-сканы раньше выпадали: fitz открывал документ, text-layer был пуст,
            # и файл ошибочно уходил в "ПДн не обнаружены". Если текста почти нет,
            # догоняем OCR по отрендеренным страницам.
            compact_pdf_text = re.sub(r"\s+", "", text_for_filter)
            if len(compact_pdf_text) < 50 and allow_ocr:
                # ДВУХФАЗНЫЙ OCR:
                # В CPU-фазе 16 воркеров не должны грузить EasyOCR в GPU (будет OOM
                # от 16×1.5 ГБ модели). Поэтому если бэкенд = "none" (CPU-пул),
                # помечаем файл как pending_ocr и уходим — GPU-воркер фазы 2
                # дообработает. Если бэкенд есть (tesseract/easyocr в одиночном
                # GPU-воркере) — OCR-им тут же.
                if _ocr_engine_name() == "none":
                    return asdict(FileResult(
                        path=str(path), format=fmt, size_bytes=size,
                        pending_ocr=True,
                        recommendation="Скан PDF без text-layer — "
                                       "отложено на GPU OCR-фазу.",
                        real_format=real_format_note,
                        elapsed_sec=time.monotonic() - t0,
                    ))

                pdf_ocr_text = extract_pdf_ocr(path)
                if pdf_ocr_text == OCR_UNAVAILABLE_SENTINEL:
                    ocr_unavailable = True
                    pdf_ocr_text = ""
                if pdf_ocr_text:
                    text_for_filter = "\n".join(
                        filter(None, [text_for_filter, pdf_ocr_text])
                    )[:MAX_TEXT_BYTES]

            found = detect_pdn(text_for_filter)
        elif suffix == ".docx":
            text_for_filter = extract_docx(path)  # уже включает embedded объекты
            found = detect_pdn(text_for_filter)
        elif suffix in (".doc", ".rtf"):
            text_for_filter = extract_doc_rtf(path)
            found = detect_pdn(text_for_filter)
        elif suffix in (".html", ".htm"):
            text_for_filter = extract_html(path)
            found = detect_pdn(text_for_filter)
        elif suffix == ".xml":
            text_for_filter = extract_xml(path)
            found = detect_pdn(text_for_filter)
        elif suffix == ".gz":
            text_for_filter = extract_gzip(path)
            found = detect_pdn(text_for_filter)
        elif suffix in (".txt", ".md"):
            text_for_filter = extract_txt(path)
            found = detect_pdn(text_for_filter)
        elif suffix in (".elf", ".exe"):
            # Бинарные исполняемые файлы — извлекаем строковые константы.
            # В реальных утилитах обработки данных строки могут содержать
            # конфигурационные ПДн: адреса серверов, email-адреса, имена.
            text_for_filter = extract_strings_binary(path)
            found = detect_pdn(text_for_filter)
        elif suffix in (".tif", ".tiff", ".jpg", ".jpeg", ".png", ".gif", ".bmp"):
            text_for_filter = extract_image_tesseract(path)
            if text_for_filter == OCR_UNAVAILABLE_SENTINEL:
                ocr_unavailable = True
                text_for_filter = ""
            found = detect_pdn(text_for_filter)

            # Дополнительно: сканируем данные после маркера EOF (DATA_AFTER_EOF).
            # Такой хвост не попадает в OCR-пайплайн — Tesseract читает только
            # пиксели изображения. Хвост может содержать текст с ПДн.
            tail_text = extract_appended_data(path, real_format_note.split()[0]
                                               if real_format_note else suffix)
            if tail_text:
                tail_found = detect_pdn(tail_text)
                _merge_findings_dict_inline(found, tail_found)
        elif suffix in (".mp4", ".avi", ".mov"):
            text_for_filter = extract_mp4_frames_ocr(path)
            if text_for_filter == VIDEO_UNAVAILABLE_SENTINEL:
                return asdict(FileResult(
                    path=str(path), format=fmt, size_bytes=size,
                    recommendation=(
                        "⚠️ Видео не обработано: OpenCV (cv2) не установлен. "
                        "Установите: pip install opencv-python. "
                        "Файл содержит видеозапись — возможны ПДн в кадрах."
                    ),
                    real_format=real_format_note,
                    elapsed_sec=time.monotonic() - t0,
                ))
            if text_for_filter == OCR_UNAVAILABLE_SENTINEL:
                ocr_unavailable = True
                text_for_filter = ""
            found = detect_pdn(text_for_filter)
        else:
            # Последний шанс: пробуем извлечь строки из неизвестного бинарника
            text_for_filter = extract_strings_binary(path)
            if text_for_filter.strip():
                found = detect_pdn(text_for_filter)
                # Помечаем что формат не определён, но попытались
                if real_format_note is None:
                    real_format_note = f"UNKNOWN (strings-scan applied)"
            else:
                return asdict(FileResult(
                    path=str(path), format=fmt, size_bytes=size,
                    recommendation="Формат не поддерживается",
                    real_format=real_format_note,
                    elapsed_sec=time.monotonic() - t0,
                ))
    except Exception as e:
        return asdict(FileResult(
            path=str(path), format=fmt, size_bytes=size,
            error=f"{type(e).__name__}: {e}",
            real_format=real_format_note,
            elapsed_sec=time.monotonic() - t0,
        ))

    # Контекстный пост-фильтр: убираем «карты» из файлов без платёжного контекста.
    # Для dict-экстракторов (CSV/Parquet/XLS/JSON) текста нет, но они и без того
    # специализированные — там «карты» скорее всего в нужной колонке.
    if text_for_filter:
        found = apply_context_filter(text_for_filter, found)

    total = sum(len(v) for v in found.values())
    cats = set(found.keys())
    uz = classify_uz(cats, total)

    # Особая рекомендация, если OCR недоступен
    if ocr_unavailable:
        recommendation = (
            "⚠️ OCR недоступен: не удалось обратиться к Tesseract. "
            "Установите https://github.com/UB-Mannheim/tesseract/wiki с поддержкой "
            "русского языка и укажите путь в переменной TESSERACT_CMD. "
            "Изображения, видео и PDF-сканы не были полноценно проверены на ПДн."
        )
    else:
        recommendation = RECOMMENDATIONS.get(uz, "")

    return asdict(FileResult(
        path=str(path),
        format=fmt,
        size_bytes=size,
        pdn_categories=found,
        total_findings=total,
        uz_level=uz,
        recommendation=recommendation,
        real_format=real_format_note,
        elapsed_sec=time.monotonic() - t0,
    ))


# ──────────────────────────────────────────────────────────────────────────────
# ОБХОД ДИРЕКТОРИЙ + ОРКЕСТРАЦИЯ
# ──────────────────────────────────────────────────────────────────────────────

EXCLUDE_DIRS = {
    ".ipynb_checkpoints", ".git", ".svn", "__pycache__",
    ".venv", "venv", "node_modules", "pdn_reports",
}


def _silence_noisy_loggers():
    """
    Отключаем warnings от PDF-библиотек и прочих «шумных» пакетов.
    На битых/старых PDF из реальных датасетов (например, ЮФУ) они сыплют
    десятки предупреждений на файл: FontBBox, Multiple definitions, /Info и т.п.
    В лог это попадает, замедляет вывод и ничего полезного не даёт.
    """
    import warnings
    warnings.filterwarnings("ignore")
    for name in (
        "pdfminer", "pdfminer.pdfparser", "pdfminer.pdfinterp",
        "pdfminer.pdfdocument", "pdfminer.cmapdb", "pdfminer.converter",
        "pdfminer.layout", "pdfminer.pdfpage", "pdfminer.psparser",
        "pypdf", "PyPDF2",
        "pdfplumber", "pdfplumber.pdf", "pdfplumber.page",
        "PIL", "PIL.TiffImagePlugin", "PIL.Image",
        "fitz",
        "chardet", "chardet.charsetprober",
        "openpyxl", "openpyxl.worksheet._reader",
        "urllib3",
    ):
        logging.getLogger(name).setLevel(logging.ERROR)


def _init_worker(
    ocr_engine: str = "auto",
    ocr_lang_tess: str = "rus+eng",
    ocr_lang_easy: Tuple[str, ...] = ("ru", "en"),
    gpu_batch_size: int = 8,
):
    """Вызывается ОДИН раз при старте каждого процесса пула.

    Важно: здесь мы сразу поднимаем OCR-бэкенд, чтобы модель EasyOCR/PaddleOCR
    (~500-1500 МБ) загрузилась в VRAM один раз на воркер, а не при каждом
    первом OCR-запросе.
    """
    global _OCR_LANG_TESS, _OCR_LANG_EASY, _GPU_BATCH_SIZE
    _OCR_LANG_TESS = ocr_lang_tess
    _OCR_LANG_EASY = ocr_lang_easy
    _GPU_BATCH_SIZE = gpu_batch_size

    _silence_noisy_loggers()
    # В воркерах не нужны INFO-логи — пишем только критичное
    logging.basicConfig(level=logging.ERROR, force=True)

    # Глушим C-уровневые stderr-сообщения MuPDF ("format error: non-page
    # object in page tree", "Failed to decode JPX image", "Memory allocation
    # failure" и т.п.). Они приходят от битых PDF (в датасете ЮФУ их десятки),
    # засоряют лог и ничего не ломают — текст уже извлечён или объявлен пустым.
    try:
        import fitz  # PyMuPDF
        fitz.TOOLS.mupdf_display_errors(False)
    except Exception:
        pass

    # Инициализируем OCR-бэкенд. Для CPU-воркеров (те, что обрабатывают
    # CSV/JSON/PDF с text-layer) это почти ничего не стоит — pytesseract
    # просто не дёргается. Для OCR-воркеров здесь грузится EasyOCR/Tesseract.
    _init_ocr_engine(ocr_engine)


def iter_files(root: Path,
               allowed_ext: Set[str],
               exclude_our_reports: bool = True) -> Iterable[Path]:
    """Рекурсивный обход. Исключаем служебные директории."""
    for dirpath, dirnames, filenames in os.walk(root, followlinks=False):
        # Фильтр имён директорий in-place — os.walk их пропустит
        dirnames[:] = [
            d for d in dirnames
            if not d.startswith(".") and d not in EXCLUDE_DIRS
        ]
        for fn in filenames:
            p = Path(dirpath) / fn
            ext = p.suffix.lower()
            if ext not in allowed_ext:
                continue
            if exclude_our_reports and ("pdn_report" in p.stem or
                                        "pii_scan_results" in p.stem):
                continue
            yield p


def load_checkpoint(jsonl_path: Path) -> Dict[str, Dict[str, Any]]:
    """Чтение уже обработанных файлов (для --resume)."""
    if not jsonl_path.exists():
        return {}
    out: Dict[str, Dict[str, Any]] = {}
    with open(jsonl_path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            try:
                rec = json.loads(line)
                out[rec["path"]] = rec
            except Exception:
                continue
    return out


def run_scan(root: Path,
             out_dir: Path,
             workers: int,
             allow_ocr: bool,
             ocr_workers: int,
             resume: bool,
             ocr_engine: str = "auto",
             ocr_lang_tess: str = "rus+eng",
             ocr_lang_easy: Tuple[str, ...] = ("ru", "en"),
             gpu_batch_size: int = 8) -> List[Dict[str, Any]]:
    out_dir.mkdir(parents=True, exist_ok=True)
    checkpoint = out_dir / "results.jsonl"

    # 1. Загружаем уже обработанные (если --resume)
    done: Dict[str, Dict[str, Any]] = {}
    if resume:
        done = load_checkpoint(checkpoint)
        log.info(f"Checkpoint: {len(done)} файлов уже обработано")
    else:
        # Начинаем заново
        if checkpoint.exists():
            checkpoint.unlink()

    # 2. Инвентаризация
    log.info(f"Инвентаризация: {root} …")
    all_files = list(iter_files(root, set(EXT_HANDLERS.keys())))
    log.info(f"Найдено файлов поддерживаемых форматов: {len(all_files)}")

    # Разделяем на OCR и не-OCR
    ocr_files: List[Path] = []
    cpu_files: List[Path] = []
    for p in all_files:
        _, kind = EXT_HANDLERS[p.suffix.lower()]
        if str(p) in done:
            continue
        if kind == "ocr":
            ocr_files.append(p)
        else:
            cpu_files.append(p)

    log.info(f"К обработке: {len(cpu_files)} текст/структура + "
             f"{len(ocr_files)} OCR (всего {len(cpu_files) + len(ocr_files)})")

    results: List[Dict[str, Any]] = list(done.values())

    # 3. CPU-bound пул — процессы.
    #    ВАЖНО: CPU-пулу ПРИНУДИТЕЛЬНО передаём ocr_engine="none".
    #    Почему: если пользователь выбрал GPU EasyOCR, то 16 CPU-воркеров,
    #    каждый со своим _init_ocr_engine("easyocr"), одновременно попытаются
    #    поднять модель (~1.5 ГБ VRAM) в GPU → 16×1.5=24 ГБ → OOM → фриз
    #    системы (это реально наблюдалось на RTX 5060 Ti 16 ГБ).
    #    PDF-сканы, для которых нужен OCR, помечаются pending_ocr=True
    #    и дообрабатываются в фазе 2 единственным GPU-воркером.
    if cpu_files:
        log.info(f"=== Фаза 1: текст/структура ({workers} процессов) ===")
        phase1_results = _run_pool(
            cpu_files, workers, allow_ocr, checkpoint,
            use_processes=True,
            ocr_engine="none",                # ← критично!
            ocr_lang_tess=ocr_lang_tess,
            ocr_lang_easy=ocr_lang_easy,
            gpu_batch_size=gpu_batch_size,
        )
        results.extend(phase1_results)

        # Собираем PDF-сканы, которые фаза 1 отложила на OCR.
        pending = [r for r in phase1_results if r.get("pending_ocr")]
        if pending:
            log.info(f"  → {len(pending)} PDF-сканов будут OCR'ены в фазе 2")
            # Убираем их из checkpoint, иначе resume пропустит.
            # Мы всё равно ПЕРЕЗАПИСЫВАЕМ эти записи в фазе 2.
            for r in pending:
                ocr_files.append(Path(r["path"]))

    # 4. OCR пул — отдельный.
    #    Для EasyOCR на GPU: ocr_workers ФОРСИРУЕМ в 1. Причины:
    #      • модель ~1.5 ГБ VRAM — N процессов × 1.5 = моментальный OOM на
    #        типичной 6-8 ГБ видеокарте (RTX 3060/4060);
    #      • даже если VRAM хватает, N процессов на ОДНОЙ GPU серализуются
    #        на уровне CUDA-драйвера → прироста нет, только overhead.
    #    Для Tesseract: оставляем как есть (N процессов, каждый своя память).
    effective_ocr_workers = ocr_workers
    if ocr_files and allow_ocr:
        _probe_engine = _probe_ocr_engine_choice(ocr_engine)
        if _probe_engine in ("easyocr", "paddleocr", "surya"):
            if ocr_workers != 1:
                log.info(f"OCR ({_probe_engine}): форсирую ocr-workers=1 "
                         f"(было {ocr_workers}) — модели не клонируются между процессами.")
            effective_ocr_workers = 1

    if ocr_files and allow_ocr:
        # Предобработка OCR-файлов: TIF→JPEG сжатие, PDF deflate
        # Уменьшает размер файлов → быстрее читаются с диска и декодируются
        import tempfile as _tempfile
        _tmp_dir = Path(_tempfile.mkdtemp(prefix="pdn_ocr_"))
        _compressed_map: Dict[Path, Path] = {}  # оригинал → сжатый
        log.info(f"Предобработка {len(ocr_files)} OCR-файлов (сжатие TIF/PDF)...")
        _comp_saved_kb = 0
        for _p in ocr_files:
            _ext = _p.suffix.lower()
            if _ext in (".tif", ".tiff", ".bmp", ".gif"):
                _cp = _compress_image_for_ocr(_p, _tmp_dir)
            elif _ext == ".pdf":
                _cp = _compress_pdf_for_ocr(_p, _tmp_dir)
            else:
                _cp = _p
            if _cp != _p:
                _compressed_map[_p] = _cp
                _comp_saved_kb += max(0, _p.stat().st_size - _cp.stat().st_size) // 1024
        if _compressed_map:
            log.info(f"  Сжато файлов: {len(_compressed_map)}, сэкономлено: {_comp_saved_kb} КБ")
            # Заменяем пути в очереди на сжатые версии
            ocr_files = [_compressed_map.get(p, p) for p in ocr_files]

        _effective_engine = _probe_ocr_engine_choice(ocr_engine)
        if _effective_engine == "surya":
            # Surya зависает в subprocess/thread на Windows — запускаем
            # прямо в главном процессе без пула.
            log.info(f"=== Фаза 2: OCR (главный процесс, движок=surya) ===")
            _init_ocr_engine("surya")
            phase2_results = _run_surya_mainprocess(
                ocr_files, allow_ocr, checkpoint,
            )
        else:
            log.info(f"=== Фаза 2: OCR ({effective_ocr_workers} процессов, движок={ocr_engine}) ===")
            phase2_results = _run_pool(
                ocr_files, effective_ocr_workers, allow_ocr, checkpoint,
                use_processes=True,
                ocr_engine=ocr_engine,
                ocr_lang_tess=ocr_lang_tess,
                ocr_lang_easy=ocr_lang_easy,
                gpu_batch_size=gpu_batch_size,
            )
        results.extend(phase2_results)
        # Удаляем временные сжатые файлы
        try:
            import shutil as _shutil
            _shutil.rmtree(_tmp_dir, ignore_errors=True)
        except Exception:
            pass
    elif ocr_files and not allow_ocr:
        # OCR отключён — просто пишем записи-заглушки
        log.info(f"OCR отключён, пропускаем {len(ocr_files)} файлов")
        for p in ocr_files:
            try:
                size = p.stat().st_size
            except OSError:
                size = 0
            rec = asdict(FileResult(
                path=str(p),
                format=EXT_HANDLERS[p.suffix.lower()][0],
                size_bytes=size,
                recommendation="OCR отключён (--no-ocr)",
            ))
            results.append(rec)
            _append_jsonl(checkpoint, rec)

    return results


def _probe_ocr_engine_choice(ocr_engine: str) -> str:
    """Определяет, какой бэкенд БУДЕТ выбран, не загружая модели.

    Нужно, чтобы понять в main-процессе, форсировать ли 1 воркер для GPU.
    Приоритет auto: paddleocr > easyocr > tesseract.
    """
    if ocr_engine == "none":
        return "none"
    if ocr_engine == "tesseract":
        return "tesseract"
    if ocr_engine == "surya":
        return "surya" if _has_surya() else "tesseract"
    if ocr_engine == "paddleocr":
        return "paddleocr" if (_detect_gpu() and _has_paddleocr()) else "tesseract"
    if ocr_engine == "easyocr":
        return "easyocr" if (_detect_gpu() and _has_easyocr()) else "tesseract"
    # auto: easyocr → tesseract
    if _detect_gpu() and _has_easyocr():
        return "easyocr"
    return "tesseract"


def _compress_image_for_ocr(src: Path, tmp_dir: Path, target_kb: int = 500) -> Path:
    """Конвертирует TIF/BMP/GIF → JPEG с целевым размером ~target_kb КБ.

    Стратегия:
      1. Открываем через Pillow, применяем _preprocess_image_for_ocr (downscale)
      2. Сохраняем как JPEG с качеством 85
      3. Если всё ещё > target_kb — снижаем quality до 60, потом 40
      4. Возвращаем путь к временному JPEG (или оригинал если уже JPEG/PNG)

    Временный файл создаётся в tmp_dir и должен быть удалён после использования.
    """
    suffix = src.suffix.lower()
    # JPEG и PNG уже оптимальны — не конвертируем
    if suffix in ('.jpg', '.jpeg', '.png'):
        return src
    try:
        from PIL import Image as PILImage
        img = PILImage.open(src)
        img = img.copy()  # avoid file handle issues
        img = _preprocess_image_for_ocr(img)

        tmp_path = tmp_dir / (src.stem + '_compressed.jpg')
        for quality in (85, 60, 40):
            img.save(tmp_path, 'JPEG', quality=quality, optimize=True)
            if tmp_path.stat().st_size <= target_kb * 1024:
                break
        log.debug(f"compress {src.name}: {src.stat().st_size//1024}КБ → {tmp_path.stat().st_size//1024}КБ")
        return tmp_path
    except Exception as e:
        log.debug(f"compress {src.name} failed: {e}")
        return src


def _compress_pdf_for_ocr(src: Path, tmp_dir: Path) -> Path:
    """Пробует сжать скан-PDF через PyMuPDF.

    Если PDF уже небольшой (<= 2МБ) или сжатие не даёт результата — возвращает
    оригинал. Временный файл создаётся в tmp_dir.
    """
    try:
        import fitz  # PyMuPDF  # type: ignore
        import os as _os
        # Глушим MuPDF stderr-спам о битых xref — ошибки некритичны,
        # fitz всё равно пытается восстановить документ.
        fitz.TOOLS.mupdf_display_errors(False)
        src_size = src.stat().st_size
        if src_size <= 2 * 1024 * 1024:  # уже <= 2МБ — не трогаем
            return src

        tmp_path = tmp_dir / (src.stem + '_compressed.pdf')
        doc = fitz.open(str(src))
        # Сохраняем с garbage collection и дефлейт-сжатием
        doc.save(
            str(tmp_path),
            garbage=4,
            deflate=True,
            deflate_images=True,
            deflate_fonts=True,
        )
        doc.close()

        new_size = tmp_path.stat().st_size
        if new_size >= src_size * 0.9:  # сжатие < 10% — не стоит
            tmp_path.unlink(missing_ok=True)
            return src

        log.debug(f"compress PDF {src.name}: {src_size//1024}КБ → {new_size//1024}КБ")
        return tmp_path
    except Exception as e:
        log.debug(f"compress PDF {src.name} failed: {e}")
        return src


def _run_surya_mainprocess(
        files: List[Path],
        allow_ocr: bool,
        checkpoint: Path,
) -> List[Dict[str, Any]]:
    """Запускает Surya OCR последовательно в главном процессе.

    Surya не работает в subprocess/thread на Windows из-за особенностей
    инициализации PyTorch CUDA. Запуск в главном процессе решает проблему.
    """
    out: List[Dict[str, Any]] = []
    total = len(files)
    start = time.monotonic()

    for i, p in enumerate(files, 1):
        try:
            rec = analyze_file(str(p), allow_ocr)
        except Exception as e:
            rec = asdict(FileResult(
                path=str(p),
                format=EXT_HANDLERS.get(p.suffix.lower(), ("UNKNOWN", ""))[0],
                size_bytes=0,
                error=f"{type(e).__name__}: {e}",
            ))
        out.append(rec)
        _append_jsonl(checkpoint, rec)

        if i % 25 == 0 or i == total:
            elapsed = time.monotonic() - start
            rate = i / elapsed if elapsed > 0 else 0
            eta = (total - i) / rate if rate > 0 else 0
            pdn = sum(1 for r in out if r["total_findings"] > 0)
            log.info(
                f"  [{i}/{total}] "
                f"скорость={rate:.1f} файл/с, "
                f"ETA={int(eta)}с, "
                f"с ПДн: {pdn}"
            )
    return out


def _run_pool(files: List[Path],
              workers: int,
              allow_ocr: bool,
              checkpoint: Path,
              use_processes: bool,
              ocr_engine: str = "auto",
              ocr_lang_tess: str = "rus+eng",
              ocr_lang_easy: Tuple[str, ...] = ("ru", "en"),
              gpu_batch_size: int = 8) -> List[Dict[str, Any]]:
    Executor = ProcessPoolExecutor if use_processes else ThreadPoolExecutor
    out: List[Dict[str, Any]] = []
    total = len(files)
    start = time.monotonic()
    # initializer запускается один раз в каждом процессе воркера.
    # Передаём ему OCR-настройки, чтобы модель EasyOCR загрузилась ровно
    # один раз (а не на каждый файл и не в main-процессе).
    exec_kwargs: Dict[str, Any] = {"max_workers": workers}
    if use_processes:
        exec_kwargs["initializer"] = _init_worker
        exec_kwargs["initargs"] = (
            ocr_engine, ocr_lang_tess, ocr_lang_easy, gpu_batch_size,
        )
    retry_files: List[Path] = []  # файлы с таймаутом — повторим в конце
    with Executor(**exec_kwargs) as pool:
        futures = {
            pool.submit(analyze_file, str(p), allow_ocr): p
            for p in files
        }
        for i, fut in enumerate(as_completed(futures), 1):
            p = futures[fut]
            try:
                rec = fut.result(timeout=FILE_TIMEOUT_SEC)
            except TimeoutError:
                # Файл завис — откладываем в конец очереди для повтора
                log.warning(f"⏱ Таймаут {FILE_TIMEOUT_SEC}с: {p.name} — добавлен в очередь повтора")
                retry_files.append(p)
                continue
            except Exception as e:
                rec = asdict(FileResult(
                    path=str(p),
                    format=EXT_HANDLERS.get(p.suffix.lower(), ("UNKNOWN", ""))[0],
                    size_bytes=0,
                    error=f"{type(e).__name__}: {e}",
                ))
            out.append(rec)
            _append_jsonl(checkpoint, rec)

            if i % 25 == 0 or i == total:
                elapsed = time.monotonic() - start
                rate = i / elapsed if elapsed > 0 else 0
                eta = (total - i) / rate if rate > 0 else 0
                pdn = sum(1 for r in out if r["total_findings"] > 0)
                log.info(
                    f"  [{i}/{total}] "
                    f"скорость={rate:.1f} файл/с, "
                    f"ETA={int(eta)}с, "
                    f"с ПДн: {pdn}"
                )

    # Повторная обработка зависших файлов с увеличенным таймаутом и логом
    if retry_files:
        log.info(f"=== Повтор {len(retry_files)} зависших файлов (таймаут {FILE_TIMEOUT_SEC*3}с) ===")
        with Executor(**exec_kwargs) as pool:
            retry_futures = {
                pool.submit(analyze_file, str(p), allow_ocr): p
                for p in retry_files
            }
            for fut in as_completed(retry_futures):
                p = retry_futures[fut]
                try:
                    rec = fut.result(timeout=FILE_TIMEOUT_SEC * 3)
                    log.info(f"  ✓ повтор OK: {p.name} — ПДн: {rec['total_findings']}")
                except Exception as e:
                    log.warning(f"  ✗ повтор FAIL: {p.name} — {type(e).__name__}: {str(e)[:80]}")
                    rec = asdict(FileResult(
                        path=str(p),
                        format=EXT_HANDLERS.get(p.suffix.lower(), ("UNKNOWN", ""))[0],
                        size_bytes=0,
                        error=f"retry {type(e).__name__}: {e}",
                    ))
                out.append(rec)
                _append_jsonl(checkpoint, rec)
    return out


def _append_jsonl(path: Path, record: Dict[str, Any]):
    with open(path, "a", encoding="utf-8") as f:
        f.write(json.dumps(record, ensure_ascii=False) + "\n")


# ──────────────────────────────────────────────────────────────────────────────
# ОТЧЁТЫ
# ──────────────────────────────────────────────────────────────────────────────

def save_csv_report(results: List[Dict[str, Any]], path: Path):
    with open(path, "w", newline="", encoding="utf-8-sig") as f:
        w = csv.DictWriter(f, fieldnames=[
            "путь", "формат_файла", "размер_байт",
            "категории_ПДн", "количество_находок",
            "УЗ", "рекомендация", "ошибка",
        ])
        w.writeheader()
        for r in results:
            cats = "; ".join(
                f"{c}({len(v)})" for c, v in (r.get("pdn_categories") or {}).items()
            )
            w.writerow({
                "путь": r["path"],
                "формат_файла": r["format"],
                "размер_байт": r["size_bytes"],
                "категории_ПДн": cats,
                "количество_находок": r["total_findings"],
                "УЗ": r["uz_level"],
                "рекомендация": r.get("recommendation", ""),
                "ошибка": r.get("error") or "",
            })


def save_json_report(results: List[Dict[str, Any]], path: Path):
    uz_count = defaultdict(int)
    for r in results:
        uz_count[r["uz_level"]] += 1
    payload = {
        "scan_date": datetime.now().isoformat(),
        "total_files": len(results),
        "files_with_pdn": sum(1 for r in results if r["total_findings"] > 0),
        "uz_summary": dict(uz_count),
        "results": results,
    }
    with open(path, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)


def _format_submission_time(ts: float) -> str:
    """Формат времени для leaderboard CSV: `sep 26 18:31`."""
    dt = datetime.fromtimestamp(ts)
    return dt.strftime("%b %d %H:%M").lower().replace(" 0", " ")


def _submission_name(path_str: str, root: Path, mode: str = "relative") -> str:
    p = Path(path_str)
    if mode == "basename":
        return p.name
    try:
        return p.relative_to(root).as_posix()
    except Exception:
        return p.name


def build_submission_rows(results: List[Dict[str, Any]], root: Path, name_mode: str = "relative") -> List[Dict[str, str]]:
    """
    Строит строки result.csv формата size,time,name.

    Берём ТОЛЬКО файлы, в которых реально найдены ПДн (total_findings > 0).
    Пустых значений быть не должно.
    """
    rows: List[Dict[str, str]] = []
    seen: Set[Tuple[str, str, str]] = set()

    for r in results:
        if int(r.get("total_findings") or 0) <= 0:
            continue

        path_str = str(r.get("path") or "").strip()
        if not path_str:
            continue

        p = Path(path_str)
        try:
            st = p.stat()
            size = int(st.st_size)
            mtime = st.st_mtime
        except OSError:
            # fallback на данные из результата, но пустые значения не допускаем
            size = int(r.get("size_bytes") or 0)
            mtime = None

        if size <= 0 or mtime is None:
            continue

        row = {
            "size": str(size),
            "time": _format_submission_time(mtime),
            "name": _submission_name(path_str, root, name_mode),
        }
        key = (row["size"], row["time"], row["name"])
        if key in seen:
            continue
        seen.add(key)
        rows.append(row)

    rows.sort(key=lambda x: x["name"].lower())
    return rows


def save_submission_csv(results: List[Dict[str, Any]], root: Path, path: Path, name_mode: str = "relative") -> int:
    rows = build_submission_rows(results, root=root, name_mode=name_mode)
    with open(path, "w", newline="", encoding="utf-8") as f:
        w = csv.DictWriter(f, fieldnames=["size", "time", "name"])
        w.writeheader()
        w.writerows(rows)
    return len(rows)


def save_markdown_report(results: List[Dict[str, Any]], path: Path):
    with_pdn = [r for r in results if r["total_findings"] > 0]
    uz_count = defaultdict(int)
    cat_count = defaultdict(int)
    for r in results:
        uz_count[r["uz_level"]] += 1
        for cat, vals in (r.get("pdn_categories") or {}).items():
            cat_count[cat] += len(vals)

    # Счётчик OCR-проблем: файлы с recommendation начинающейся с "⚠️ OCR"
    ocr_unavailable_cnt = sum(
        1 for r in results
        if r.get("recommendation", "").startswith("⚠️ OCR недоступен")
    )

    # Файлы с подменённым форматом (расширение ≠ реальный тип)
    spoofed_files = [r for r in results if r.get("real_format")]

    L = [
        "# Отчёт сканирования персональных данных (152-ФЗ)",
        "",
        f"**Дата:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  ",
        f"**Всего файлов:** {len(results)}  ",
        f"**С ПДн:** {len(with_pdn)}  ",
        f"**Без ПДн:** {len(results) - len(with_pdn)}",
        "",
    ]

    if ocr_unavailable_cnt > 0:
        L += [
            "> ⚠️ **Предупреждение об OCR**",
            f"> {ocr_unavailable_cnt} изображений/видео не были проверены — Tesseract "
            "не установлен. Среди них могут быть сканы паспортов / СНИЛС / иных "
            "документов. Установите Tesseract и запустите сканер повторно, чтобы "
            "не пропустить ПДн в сканированных файлах.",
            "",
        ]

    if spoofed_files:
        L += [
            "> 🔍 **Файлы с несоответствием расширения и реального формата**",
            f"> Обнаружено **{len(spoofed_files)}** файлов, чей реальный формат "            "отличается от расширения (например, .pdf → HTML, .docx → OLE2). "            "Такие файлы были проанализированы по реальному формату. "            "Это может свидетельствовать об ошибке при сохранении файла "            "или намеренном переименовании.",
            "",
        ]

    L += [
        "## Сводка по уровням защищённости",
        "",
        "| УЗ | Файлов | Описание |",
        "|----|-------|----------|",
        f"| УЗ-1 | {uz_count['УЗ-1']} | Биометрия / спец. категории |",
        f"| УЗ-2 | {uz_count['УЗ-2']} | Платёжные данные / много гос. ID |",
        f"| УЗ-3 | {uz_count['УЗ-3']} | Гос. идентификаторы / много обычных ПДн |",
        f"| УЗ-4 | {uz_count['УЗ-4']} | Обычные ПДн в небольшом объёме |",
        f"| — | {uz_count['—']} | ПДн не обнаружены |",
        "",
        "## Категории ПДн (суммарно по всем файлам)",
        "",
        "| Категория | Находок |",
        "|-----------|--------:|",
    ]
    for cat, cnt in sorted(cat_count.items(), key=lambda x: -x[1]):
        L.append(f"| {cat} | {cnt} |")

    L += ["", "## Файлы с обнаруженными ПДн", ""]
    order = {"УЗ-1": 0, "УЗ-2": 1, "УЗ-3": 2, "УЗ-4": 3, "—": 4}
    with_pdn.sort(key=lambda r: (order.get(r["uz_level"], 5), -r["total_findings"]))

    # Ограничим детализацию первыми 300 файлами, чтобы md не стал гигантским
    for r in with_pdn[:300]:
        L.append(f"### `{r['path']}`")
        fmt_str = r['format']
        if r.get('real_format'):
            fmt_str += f" ⚠️ реальный: {r['real_format']}"
        L.append(f"- Формат: {fmt_str} · Размер: {r['size_bytes']:,} байт")
        L.append(f"- **УЗ: {r['uz_level']}** · Находок: {r['total_findings']}")
        L.append("- Категории:")
        for cat, vals in (r.get("pdn_categories") or {}).items():
            sample = ", ".join(f"`{v}`" for v in vals[:3])
            more = f" (+{len(vals) - 3})" if len(vals) > 3 else ""
            L.append(f"  - {cat}: {len(vals)} экз — {sample}{more}")
        L.append(f"- Рекомендация: {r.get('recommendation', '')}")
        L.append("")

    if len(with_pdn) > 300:
        L.append(f"_… и ещё {len(with_pdn) - 300} файлов — см. CSV/JSON отчёт._")

    # Секция: файлы с несоответствием расширения и реального формата
    if spoofed_files:
        L += ["", "## Файлы с подменённым форматом", ""]
        L.append(
            "Эти файлы были обработаны **по реальному формату** (magic bytes), "
            "а не по расширению. Возможные причины: экспорт из Word/Excel в HTML "
            "без смены расширения, ошибка при сохранении, намеренное переименование."
        )
        L.append("")
        L.append("| Файл | Заявленное расширение | Реальный формат |")
        L.append("|------|-----------------------|-----------------|")
        for r in spoofed_files:
            rf = r.get("real_format") or ""
            import re as _re
            m = _re.search(r"заявлен как \'([^\']+)\'", rf)
            ext_claimed = m.group(1) if m else "?"
            real_label = rf.split(" [")[0]
            L.append(f"| `{r['path']}` | `{ext_claimed}` | {real_label} |")
        L.append("")

    errs = [r for r in results if r.get("error")]
    if errs:
        L += ["", "## Ошибки обработки", ""]
        for r in errs[:100]:
            L.append(f"- `{r['path']}` — {r['error']}")
        if len(errs) > 100:
            L.append(f"_… и ещё {len(errs) - 100}_")

    path.write_text("\n".join(L), encoding="utf-8")


# ──────────────────────────────────────────────────────────────────────────────
# MAIN
# ──────────────────────────────────────────────────────────────────────────────

def main() -> int:
    ap = argparse.ArgumentParser(
        description="Сканер персональных данных (152-ФЗ) с опциональным GPU OCR",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""\
Примеры:
  # Автодетект: paddleocr → easyocr → tesseract (по убыванию скорости)
  python pdn_scanner.py /path/to/dataset

  # Форсировать PaddleOCR (самый быстрый на GPU)
  python pdn_scanner.py /path/to/dataset --ocr-engine paddleocr

  # EasyOCR (если paddlepaddle не ставится из-за CUDA-версии)
  python pdn_scanner.py /path/to/dataset --ocr-engine easyocr

  # Tesseract на CPU (без GPU)
  python pdn_scanner.py /path/to/dataset --ocr-engine tesseract --ocr-workers 3

  # Большой батч для GPU с 12+ ГБ VRAM (RTX 4070/4080/5060Ti, A100)
  python pdn_scanner.py /path/to/dataset --gpu-batch-size 16

  # Только текст, без OCR
  python pdn_scanner.py /path/to/dataset --no-ocr

УСТАНОВКА PaddleOCR + GPU (рекомендуется, быстрейший вариант):
  # CUDA 12.6 (для RTX 30/40/50 серии):
  pip install paddlepaddle-gpu==3.0.0 -i https://www.paddlepaddle.org.cn/packages/stable/cu126/
  # CUDA 11.8 (для старых карт):
  # pip install paddlepaddle-gpu==3.0.0 -i https://www.paddlepaddle.org.cn/packages/stable/cu118/
  pip install paddleocr

УСТАНОВКА EasyOCR + GPU (fallback, если paddlepaddle не ставится):
  pip install torch torchvision --index-url https://download.pytorch.org/whl/cu121
  pip install easyocr
""",
    )
    ap.add_argument("directory", help="Путь к директории для сканирования")
    ap.add_argument("--out-dir", default="pdn_reports",
                    help="Директория для отчётов (по умолчанию: pdn_reports)")
    ap.add_argument("--workers", type=int, default=max(1, os.cpu_count() or 4),
                    help="Процессов для текст/структура (по умолч: CPU cores)")
    ap.add_argument("--ocr-workers", type=int, default=2,
                    help="Процессов для OCR (CPU Tesseract; для GPU EasyOCR "
                         "форсируется 1)")
    ap.add_argument("--no-ocr", action="store_true",
                    help="Пропустить изображения и видео")
    ap.add_argument("--ocr-engine",
                    default="auto",
                    choices=["auto", "surya", "paddleocr", "easyocr", "tesseract", "none"],
                    help="OCR-движок. auto=paddleocr→easyocr→tesseract. "
                         "paddleocr обычно в 3-10× быстрее easyocr. "
                         "none=отключить OCR целиком.")
    ap.add_argument("--gpu-batch-size", type=int, default=8,
                    help="Размер батча распознавания для EasyOCR (по умолч: 8). "
                         "12+ ГБ VRAM — поднимайте до 16-32 для ускорения.")
    ap.add_argument("--ocr-lang-tess", default="rus+eng",
                    help="Языки Tesseract (по умолч: rus+eng)")
    ap.add_argument("--ocr-lang-easy", default="ru,en",
                    help="Языки EasyOCR через запятую (по умолч: ru,en)")
    ap.add_argument("--resume", action="store_true",
                    help="Продолжить с results.jsonl (если был прерван)")
    ap.add_argument("--log-level", default="INFO",
                    choices=["DEBUG", "INFO", "WARNING", "ERROR"])
    ap.add_argument("--result-csv", default="result.csv",
                    help="Итоговый CSV для сабмита leaderboard (по умолчанию: result.csv)")
    ap.add_argument("--submission-name-mode", default="basename",
                    choices=["relative", "basename"],
                    help="Что писать в колонку name: относительный путь или только имя файла")
    args = ap.parse_args()

    logging.getLogger().setLevel(getattr(logging, args.log_level))

    # Глушим warnings шумных библиотек и в главном процессе
    _silence_noisy_loggers()

    root = Path(args.directory).resolve()
    if not root.exists():
        log.error(f"Директория не найдена: {root}")
        return 1
    if not root.is_dir():
        log.error(f"Не директория: {root}")
        return 1

    out_dir = Path(args.out_dir).resolve()
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")

    # Определяем, какой OCR-бэкенд РЕАЛЬНО будет выбран — для баннера и для
    # автокоррекции числа OCR-воркеров.
    effective_engine = _probe_ocr_engine_choice(args.ocr_engine)
    ocr_lang_easy_tuple = tuple(
        x.strip() for x in args.ocr_lang_easy.split(",") if x.strip()
    )

    log.info("=" * 70)
    log.info(f"ПДн-СКАНЕР")
    log.info(f"  директория     : {root}")
    log.info(f"  отчёты         : {out_dir}")
    log.info(f"  workers (CPU)  : {args.workers}")
    log.info(f"  workers (OCR)  : {args.ocr_workers}"
             f"{' → 1 (GPU)' if effective_engine == 'easyocr' else ''}")
    log.info(f"  OCR            : {'ВЫКЛ' if args.no_ocr else 'ВКЛ'}")
    log.info(f"  OCR-движок     : {args.ocr_engine} → {effective_engine}")
    if effective_engine == "easyocr":
        log.info(f"  GPU batch size : {args.gpu_batch_size}")
        log.info(f"  OCR languages  : {ocr_lang_easy_tuple}")
    elif effective_engine == "tesseract":
        log.info(f"  OCR languages  : {args.ocr_lang_tess}")
    log.info(f"  resume         : {args.resume}")
    log.info(f"  result.csv     : {Path(args.result_csv).resolve()}")
    log.info(f"  name mode      : {args.submission_name_mode}")
    log.info("=" * 70)

    # Проверка GPU (только если выбран easyocr/paddleocr или auto) —
    # максимально информативный баннер, чтобы пользователь сразу понимал,
    # что происходит.
    if not args.no_ocr and args.ocr_engine in ("auto", "easyocr", "paddleocr"):
        try:
            import torch  # type: ignore
            if torch.cuda.is_available():
                try:
                    name = torch.cuda.get_device_name(0)
                    vram_gb = torch.cuda.get_device_properties(0).total_memory / (1024**3)
                    log.info(f"✓ CUDA GPU обнаружен: {name} ({vram_gb:.1f} ГБ VRAM)")
                except Exception:
                    log.info("✓ CUDA GPU обнаружен")
            else:
                log.info("CUDA GPU не обнаружен — GPU OCR-движки недоступны")
        except ImportError:
            log.info("PyTorch не установлен (нужен для GPU-детекта)")

        if effective_engine == "paddleocr":
            log.info("✓ PaddleOCR установлен — будет использован (fastest on GPU)")
        elif effective_engine == "easyocr":
            if _has_paddleocr():
                log.info("✓ EasyOCR выбран (PaddleOCR тоже установлен — "
                         "можно ускориться через --ocr-engine paddleocr)")
            else:
                log.info("✓ EasyOCR установлен — GPU OCR будет использован")
                log.info("  ⚡ Для ещё большей скорости попробуйте PaddleOCR:")
                log.info("     pip install paddlepaddle-gpu paddleocr")
        elif effective_engine == "tesseract":
            # Спец-предупреждения если пользователь просил GPU, но нет
            if args.ocr_engine == "paddleocr":
                log.warning("⚠️  Запрошен --ocr-engine=paddleocr, но PaddleOCR "
                            "не установлен или CUDA недоступна. Откат на Tesseract.")
            elif args.ocr_engine == "easyocr":
                log.warning("⚠️  Запрошен --ocr-engine=easyocr, но easyocr "
                            "не установлен или CUDA недоступна. Откат на Tesseract.")

    # Проверка Tesseract — нужна только если реально пойдём через него
    if not args.no_ocr and effective_engine == "tesseract":
        if tesseract_available():
            log.info("✓ Tesseract OCR доступен — изображения будут обработаны")
        else:
            log.warning("=" * 70)
            log.warning("⚠️  ВНИМАНИЕ: Tesseract НЕ найден!")
            log.warning("  Все изображения (TIF/JPEG/PNG/GIF) и видео будут")
            log.warning("  помечены как «OCR недоступен». В датасетах с TIF-сканами")
            log.warning("  паспортов это означает, что ПДн на сканах НЕ обнаружены.")
            log.warning("")
            log.warning("  Варианты:")
            log.warning("    1) Установить EasyOCR (если есть NVIDIA GPU — лучший выбор):")
            log.warning("       pip install easyocr torch")
            log.warning("    2) Установить Tesseract:")
            log.warning("       https://github.com/UB-Mannheim/tesseract/wiki")
            log.warning("       (обязательно с пакетом 'Russian')")
            log.warning("    3) Запустить с --no-ocr — честно пометить все")
            log.warning("       изображения/видео как непроверенные.")
            log.warning("=" * 70)

    # Проверка OpenCV — нужен для обработки видео (MP4/AVI/MOV).
    # dataset4 содержит видеозаписи сканов документов — без cv2 они не обработаются.
    if not args.no_ocr:
        if cv2_available():
            log.info("✓ OpenCV (cv2) доступен — видеофайлы будут обработаны")
        else:
            log.warning("=" * 70)
            log.warning("⚠️  ВНИМАНИЕ: OpenCV (cv2) НЕ найден!")
            log.warning("  Видеофайлы (MP4/AVI/MOV) не будут обработаны.")
            log.warning("  В датасете dataset4 содержатся видео-сканы документов —")
            log.warning("  без cv2 ПДн в них обнаружены НЕ будут.")
            log.warning("")
            log.warning("  Установка: pip install opencv-python")
            log.warning("=" * 70)

    t0 = time.monotonic()
    try:
        results = run_scan(
            root=root,
            out_dir=out_dir,
            workers=args.workers,
            allow_ocr=not args.no_ocr,
            ocr_workers=args.ocr_workers,
            resume=args.resume,
            ocr_engine=args.ocr_engine,
            ocr_lang_tess=args.ocr_lang_tess,
            ocr_lang_easy=ocr_lang_easy_tuple,
            gpu_batch_size=args.gpu_batch_size,
        )
    except KeyboardInterrupt:
        log.warning("Прервано пользователем. Частичные результаты — "
                    f"в {out_dir / 'results.jsonl'}. "
                    f"Запустите с --resume для продолжения.")
        return 130

    elapsed = time.monotonic() - t0

    # Дедупликация по пути (resume мог дать дубли)
    dedup: Dict[str, Dict[str, Any]] = {}
    for r in results:
        dedup[r["path"]] = r
    results = list(dedup.values())

    log.info("=" * 70)
    log.info(f"Обработка завершена за {int(elapsed // 60)} мин {int(elapsed % 60)} сек")
    log.info(f"Всего файлов: {len(results)}")
    pdn_files = [r for r in results if r["total_findings"] > 0]
    log.info(f"Файлов с ПДн: {len(pdn_files)}")
    err_files = [r for r in results if r.get("error")]
    if err_files:
        log.warning(f"Ошибок: {len(err_files)} (см. отчёты)")

    # Сохраняем финальные отчёты
    save_csv_report(results, out_dir / f"pdn_report_{ts}.csv")
    log.info(f"CSV:      {out_dir / f'pdn_report_{ts}.csv'}")
    save_json_report(results, out_dir / f"pdn_report_{ts}.json")
    log.info(f"JSON:     {out_dir / f'pdn_report_{ts}.json'}")
    save_markdown_report(results, out_dir / f"pdn_report_{ts}.md")
    log.info(f"Markdown: {out_dir / f'pdn_report_{ts}.md'}")

    submission_path = Path(args.result_csv).resolve()
    submission_count = save_submission_csv(
        results,
        root=root,
        path=submission_path,
        name_mode=args.submission_name_mode,
    )
    log.info(f"result.csv: {submission_path} (строк: {submission_count})")

    # Критические находки
    uz1 = [r for r in results if r["uz_level"] == "УЗ-1"]
    if uz1:
        log.warning(f"🚨 УЗ-1 (биометрия/спец. категории): {len(uz1)} файлов")
        for r in uz1[:5]:
            log.warning(f"   → {r['path']}")
    # Предупреждение о файлах с подменённым форматом
    spoofed = [r for r in results if r.get("real_format")]
    if spoofed:
        log.warning(f"🔍 Файлы с несоответствием расширения и реального формата: {len(spoofed)}")
        for r in spoofed[:10]:
            log.warning(f"   → {r['path']}  ({r.get('real_format')})")
        if len(spoofed) > 10:
            log.warning(f"   … и ещё {len(spoofed) - 10} — см. Markdown-отчёт")
    log.info("=" * 70)
    return 0


if __name__ == "__main__":
    # На Windows нужен 'spawn' — он и так по умолчанию, но на всякий случай:
    try:
        import multiprocessing as mp
        if mp.get_start_method(allow_none=True) is None:
            mp.set_start_method("spawn")
    except Exception:
        pass
    raise SystemExit(main())